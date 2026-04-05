# ════════════════════════════════════════════════════════════════════════════════
# SOMO AI v4.0 — COMPLETE APPLICATION
# ════════════════════════════════════════════════════════════════════════════════
# Author  : Usmonov Sodiq
# Date    : 2026-03-25
# Version : 4.0  (Secure · Mobile-Ready · All Bugs Fixed)
# Lines   : 2200+
# ════════════════════════════════════════════════════════════════════════════════

import streamlit as st
import pandas as pd
import gspread
import json
import time
import io
import re
import os
from datetime import datetime
from functools import lru_cache
from typing import Tuple, Optional, List, Dict, Any

# ════════════════════════════════════════════════════════════════════════════════
# OPTIONAL IMPORTS
# ════════════════════════════════════════════════════════════════════════════════

try:
    from oauth2client.service_account import ServiceAccountCredentials
    HAS_OAUTH = True
except ImportError:
    HAS_OAUTH = False

try:
    import bcrypt
    HAS_BCRYPT = True
except ImportError:
    HAS_BCRYPT = False

try:
    from streamlit_cookies_manager import EncryptedCookieManager
    HAS_COOKIES = True
except ImportError:
    HAS_COOKIES = False

try:
    from mammoth import extract_raw_text as mammoth_extract
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from groq import Groq
    HAS_GROQ = True
except ImportError:
    HAS_GROQ = False

try:
    import google.generativeai as genai
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

try:
    import cohere
    HAS_COHERE = True
    # v2 SDK (cohere>=5.0) da ClientV2 bor
    _COHERE_V2 = hasattr(cohere, "ClientV2")
except ImportError:
    HAS_COHERE = False
    _COHERE_V2 = False

try:
    from mistralai import Mistral
    HAS_MISTRAL = True
except ImportError:
    HAS_MISTRAL = False

# ════════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ════════════════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="Somo AI | Ultra Pro Max",
    page_icon="🌌",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ════════════════════════════════════════════════════════════════════════════════
# MASTER CSS
# ════════════════════════════════════════════════════════════════════════════════

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;700;800&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

:root {
  --bg-0:    #05050f;
  --bg-1:    #090916;
  --bg-2:    #0d0d1e;
  --bg-card: #0f0f22;
  --border:  rgba(100,108,255,0.18);
  --border-h:rgba(100,108,255,0.55);
  --accent:  #646cff;
  --text-1:  #f0f0ff;
  --text-2:  #a0a0c0;
  --text-3:  #50506a;
  --fhead:   'Syne', sans-serif;
  --fbody:   'Inter', sans-serif;
  --fmono:   'JetBrains Mono', monospace;
  --r:       16px;
  --r-sm:    10px;
}

html, body, .stApp {
  font-family: var(--fbody) !important;
  background: var(--bg-0) !important;
  color: var(--text-1) !important;
}

/* ── Hide chrome ── */
header[data-testid="stHeader"],
#MainMenu, footer,
[data-testid="stSidebarNav"],
.st-emotion-cache-1vt458p { display: none !important; }

/* ── Main padding ── */
.main .block-container,
[data-testid="stMainBlockContainer"] {
  padding: 1.5rem 1rem 5rem !important;
  max-width: 100% !important;
}
@media (min-width: 768px) {
  .main .block-container,
  [data-testid="stMainBlockContainer"] {
    padding: 1.5rem 2rem 5rem !important;
    max-width: 1100px !important;
  }
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: var(--bg-1); }
::-webkit-scrollbar-thumb { background: #2a2a55; border-radius: 99px; }
::-webkit-scrollbar-thumb:hover { background: var(--accent); }

/* ════ SIDEBAR ════ */
[data-testid="stSidebar"] {
  background: linear-gradient(180deg,#07071a 0%,#09091e 100%) !important;
  border-right: 1px solid rgba(100,108,255,0.15) !important;
  width: 260px !important;
  display: block !important;
  z-index: 999 !important;
}
[data-testid="collapsedControl"] {
  display: flex !important;
  z-index: 1000 !important;
  color: #818cf8 !important;
}
[data-testid="stSidebar"] > div:first-child { padding: 0 !important; }

[data-testid="stSidebar"] .stButton > button {
  background: transparent !important;
  color: var(--text-2) !important;
  border: none !important;
  border-radius: var(--r-sm) !important;
  font-weight: 500 !important;
  font-size: 13px !important;
  padding: 9px 14px !important;
  width: 100% !important;
  text-align: left !important;
  transition: all 0.2s !important;
  margin: 1px 0 !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
  background: rgba(100,108,255,0.1) !important;
  color: #c7d2fe !important;
  transform: translateX(3px) !important;
}
[data-testid="stSidebar"] .stButton > button[kind="primary"] {
  background: linear-gradient(135deg,rgba(100,108,255,0.2),rgba(167,139,250,0.12)) !important;
  color: #c7d2fe !important;
  border: 1px solid rgba(100,108,255,0.3) !important;
  font-weight: 700 !important;
}
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stSlider label,
[data-testid="stSidebar"] [data-testid="stWidgetLabel"] {
  color: var(--text-2) !important;
  font-size: 12px !important;
  font-family: var(--fmono) !important;
}
[data-testid="stSidebar"] div[data-baseweb="select"] > div {
  background: rgba(100,108,255,0.06) !important;
  border-color: rgba(100,108,255,0.2) !important;
  color: var(--text-1) !important;
  border-radius: var(--r-sm) !important;
}
[data-testid="stSidebar"] .stDownloadButton > button {
  background: rgba(52,211,153,0.1) !important;
  border: 1px solid rgba(52,211,153,0.25) !important;
  color: #6ee7b7 !important;
  font-size: 12px !important;
  padding: 8px 12px !important;
}

/* ════ HERO ════ */
.somo-hero {
  position: relative; overflow: hidden;
  border-radius: 24px; padding: 52px 48px; margin-bottom: 32px;
  background: var(--bg-1);
  border: 1px solid var(--border);
  box-shadow: 0 4px 40px rgba(0,0,0,0.6),0 0 40px rgba(100,108,255,0.12);
}
.somo-hero::before {
  content:''; position:absolute; top:-80px; left:-80px;
  width:500px; height:500px; border-radius:50%;
  background:radial-gradient(circle,rgba(100,108,255,0.14) 0%,transparent 60%);
  animation: porb 10s ease-in-out infinite;
}
.somo-hero::after {
  content:''; position:absolute; bottom:-60px; right:-40px;
  width:400px; height:400px; border-radius:50%;
  background:radial-gradient(circle,rgba(244,114,182,0.1) 0%,transparent 60%);
  animation: porb 10s ease-in-out 5s infinite;
}
@keyframes porb {
  0%,100%{transform:scale(1) translate(0,0)}
  50%{transform:scale(1.2) translate(-15px,-15px)}
}
.grid-dots {
  position:absolute; inset:0;
  background-image:radial-gradient(circle,rgba(100,108,255,0.18) 1px,transparent 1px);
  background-size:32px 32px; opacity:0.25;
}
.somo-hero-content { position:relative; z-index:2; }
.somo-hero h1 {
  font-family: var(--fhead) !important;
  font-size: clamp(28px,5vw,52px);
  font-weight: 800; line-height: 1.1;
  letter-spacing: -2px; color: var(--text-1); margin-bottom: 12px;
}
.somo-hero .subtitle {
  font-size: 15px; color: var(--text-2);
  max-width: 520px; line-height: 1.65; margin-bottom: 20px;
}
.hero-badges { display:flex; flex-wrap:wrap; gap:8px; }
.hero-badge {
  background:rgba(255,255,255,0.06); border:1px solid rgba(255,255,255,0.12);
  color:rgba(255,255,255,0.65); padding:5px 13px; border-radius:99px;
  font-size:11.5px; font-weight:500; font-family:var(--fmono);
}

/* ════ GRADIENT TEXT ════ */
.g-text {
  background: linear-gradient(90deg,#818cf8,#c084fc,#f472b6,#818cf8);
  background-size:300%;
  -webkit-background-clip:text; -webkit-text-fill-color:transparent;
  background-clip:text; animation:gshift 5s ease infinite;
}
@keyframes gshift { 0%,100%{background-position:0%} 50%{background-position:100%} }

/* ════ CARDS ════ */
.cards-grid {
  display:grid;
  grid-template-columns:repeat(auto-fill,minmax(180px,1fr));
  gap:14px; margin-bottom:28px;
}
@media(max-width:480px){ .cards-grid{grid-template-columns:repeat(2,1fr);gap:8px;} }
.somo-card {
  background:var(--bg-card); border:1px solid var(--border);
  border-radius:var(--r); padding:26px 16px; text-align:center;
  transition:all 0.3s cubic-bezier(.4,0,.2,1); cursor:pointer;
  position:relative; overflow:hidden;
}
.somo-card:hover { transform:translateY(-6px); border-color:var(--border-h); box-shadow:0 20px 50px rgba(0,0,0,0.5),0 0 25px rgba(100,108,255,0.15); }
.card-icon  { font-size:32px; margin-bottom:12px; display:block; }
.card-title { font-size:13.5px; font-weight:700; color:var(--text-1); margin-bottom:5px; font-family:var(--fhead); }
.card-desc  { font-size:11px; color:var(--text-3); line-height:1.55; }

/* ════ STATS ════ */
.stat-row { display:grid; grid-template-columns:repeat(auto-fill,minmax(120px,1fr)); gap:10px; margin-bottom:24px; }
@media(max-width:640px){ .stat-row{grid-template-columns:repeat(2,1fr);} }
.stat-box { background:var(--bg-card); border:1px solid var(--border); border-radius:14px; padding:16px 12px; text-align:center; transition:border-color .3s; }
.stat-box:hover { border-color:var(--border-h); }
.stat-icon { font-size:20px; margin-bottom:6px; }
.stat-val  { font-size:26px; font-weight:900; color:var(--text-1); font-family:var(--fhead); }
.stat-lbl  { font-size:9.5px; font-weight:700; color:var(--text-3); margin-top:4px; text-transform:uppercase; letter-spacing:1.5px; font-family:var(--fmono); }

/* ════ DIVIDER / LABELS ════ */
.somo-divider {
  height:1px; background:linear-gradient(90deg,transparent,var(--border),transparent);
  margin:24px 0; border:none;
}
.section-label {
  font-size:10px; font-weight:700; letter-spacing:2.5px;
  text-transform:uppercase; color:var(--accent); margin-bottom:8px;
  font-family:var(--fmono); display:flex; align-items:center; gap:8px;
}
.section-label::before { content:''; display:inline-block; width:14px; height:1px; background:var(--accent); }
.section-title { font-size:22px; font-weight:800; color:var(--text-1); margin-bottom:6px; font-family:var(--fhead); letter-spacing:-0.5px; }

/* ════ CHAT ════ */
.stChatMessage {
  background:var(--bg-card) !important; border:1px solid var(--border) !important;
  border-radius:var(--r) !important; padding:14px 18px !important; margin:6px 0 !important;
  color:var(--text-1) !important;
}
.stChatMessage p,.stChatMessage span,.stChatMessage li { color:var(--text-1) !important; }
.stChatMessage code { background:rgba(100,108,255,0.12) !important; color:#a5b4fc !important; border-radius:4px; padding:1px 6px; font-family:var(--fmono) !important; }
.stChatMessage pre  { background:#04040f !important; border:1px solid var(--border) !important; border-radius:var(--r-sm) !important; }

[data-testid="stChatInputContainer"],
[data-testid="stChatInputContainer"] > div,
[data-testid="stChatInputContainer"] > div > div {
  background: rgba(13,13,30,0.9) !important;
  border: 1px solid var(--border) !important;
  border-radius: 14px !important;
}
[data-testid="stChatInputContainer"] textarea {
  background: transparent !important; color: var(--text-1) !important;
  font-family: var(--fbody) !important; font-size: 14px !important;
  caret-color: var(--accent) !important;
}
[data-testid="stChatInputContainer"] button {
  background: linear-gradient(135deg,#4f46e5,#7c3aed) !important;
  border: none !important; border-radius: var(--r-sm) !important; color: white !important;
}
div[data-testid="stBottom"],
div[data-testid="stBottom"] > div,
div[data-testid="stBottom"] > div > div {
  background: rgba(5,5,15,0.97) !important;
  backdrop-filter: blur(20px) !important;
  border-top: 1px solid var(--border) !important;
}

/* ════ INPUTS ════ */
.stTextInput input, .stTextArea textarea, div[data-baseweb="input"] input {
  background: var(--bg-2) !important; color: var(--text-1) !important;
  border: 1px solid var(--border) !important; border-radius: var(--r-sm) !important;
  font-family: var(--fbody) !important; font-size: 14px !important;
}
.stTextInput input:focus,.stTextArea textarea:focus {
  border-color: var(--accent) !important;
  box-shadow: 0 0 0 3px rgba(100,108,255,0.1) !important;
}
.stTextInput input::placeholder,.stTextArea textarea::placeholder { color: var(--text-3) !important; }
.stTextInput label,.stTextArea label,.stSelectbox label,
[data-testid="stWidgetLabel"] { color:var(--text-2) !important; font-size:13px !important; font-weight:600 !important; font-family:var(--fmono) !important; }

div[data-baseweb="select"] > div {
  background: var(--bg-2) !important; border-color: var(--border) !important;
  border-radius: var(--r-sm) !important; color: var(--text-1) !important;
}
div[data-baseweb="popover"] {
  background: var(--bg-2) !important; border: 1px solid var(--border) !important;
  border-radius: var(--r) !important; box-shadow: 0 4px 40px rgba(0,0,0,0.6) !important;
}
div[data-baseweb="popover"] li { color:var(--text-1) !important; }
div[data-baseweb="popover"] li:hover { background:rgba(100,108,255,0.1) !important; }

[data-testid="stFileUploader"] {
  background: rgba(100,108,255,0.04) !important;
  border: 2px dashed rgba(100,108,255,0.2) !important;
  border-radius: var(--r) !important; padding: 16px !important;
}
[data-testid="stFileUploader"]:hover { border-color:rgba(100,108,255,0.4) !important; }
[data-testid="stFileUploader"] * { color: var(--text-2) !important; }
[data-testid="stFileUploaderDropzone"] {
  background: rgba(13,13,30,0.6) !important;
  border: 1px dashed rgba(100,108,255,0.25) !important;
  border-radius: 12px !important;
}
[data-testid="stFileUploaderDropzone"] button {
  background: rgba(100,108,255,0.1) !important;
  color: #818cf8 !important;
  border: 1px solid rgba(100,108,255,0.25) !important;
  border-radius: 8px !important;
}

/* ════ BUTTONS ════ */
.stButton > button {
  background: rgba(100,108,255,0.07) !important; color: #a5b4fc !important;
  border: 1px solid var(--border) !important; border-radius: var(--r-sm) !important;
  font-family: var(--fbody) !important; font-weight: 600 !important;
  font-size: 13px !important; padding: 9px 18px !important; transition: all 0.2s !important;
}
.stButton > button:hover {
  background: rgba(100,108,255,0.15) !important; color: #c7d2fe !important;
  border-color: var(--border-h) !important; transform: translateY(-2px) !important;
  box-shadow: 0 6px 20px rgba(100,108,255,0.18) !important;
}
.stButton > button[kind="primary"] {
  background: linear-gradient(135deg,#4f46e5,#7c3aed) !important;
  color: white !important; border-color: transparent !important;
  box-shadow: 0 0 20px rgba(100,108,255,0.28) !important;
}
.stButton > button[kind="primary"]:hover {
  background: linear-gradient(135deg,#4338ca,#6d28d9) !important;
  box-shadow: 0 6px 28px rgba(100,108,255,0.45) !important; transform: translateY(-3px) !important;
}
.stDownloadButton > button {
  background: linear-gradient(135deg,#059669,#10b981) !important;
  color: white !important; border: none !important;
  border-radius: var(--r-sm) !important; font-weight: 700 !important;
  box-shadow: 0 4px 20px rgba(16,185,129,0.3) !important; transition: all 0.25s !important;
}
.stDownloadButton > button:hover { transform:translateY(-3px) !important; box-shadow:0 8px 28px rgba(16,185,129,0.45) !important; }

/* ════ TABS ════ */
.stTabs [data-baseweb="tab-list"] { background:transparent !important; border-bottom:1px solid var(--border) !important; gap:4px !important; }
.stTabs [data-baseweb="tab"] { background:transparent !important; color:var(--text-3) !important; font-weight:600 !important; font-size:13px !important; padding:9px 18px !important; }
.stTabs [data-baseweb="tab"]:hover { color:#a5b4fc !important; background:rgba(100,108,255,0.07) !important; }
.stTabs [aria-selected="true"][data-baseweb="tab"] { color:#818cf8 !important; border-bottom:2px solid var(--accent) !important; }
.stTabs [data-baseweb="tab-panel"] { background:transparent !important; padding:18px 0 !important; }

/* ════ FORMS ════ */
[data-testid="stForm"] {
  background: linear-gradient(145deg,var(--bg-card),var(--bg-2)) !important;
  border: 1px solid var(--border) !important; border-radius: var(--r) !important; padding: 24px !important;
}
[data-testid="stFormSubmitButton"] > button {
  background: linear-gradient(135deg,#4f46e5,#7c3aed) !important;
  color: white !important; border: none !important; font-weight: 700 !important;
}

/* ════ ALERTS ════ */
.stSuccess > div { background:rgba(52,211,153,0.08) !important; border:1px solid rgba(52,211,153,0.25) !important; color:#6ee7b7 !important; border-radius:var(--r-sm) !important; }
.stWarning > div { background:rgba(251,191,36,0.08) !important; border:1px solid rgba(251,191,36,0.25) !important; color:#fcd34d !important; border-radius:var(--r-sm) !important; }
.stError   > div { background:rgba(244,114,182,0.08) !important; border:1px solid rgba(244,114,182,0.25) !important; color:#fca5a5 !important; border-radius:var(--r-sm) !important; }
.stInfo    > div { background:rgba(100,108,255,0.08) !important; border:1px solid rgba(100,108,255,0.25) !important; color:#a5b4fc !important; border-radius:var(--r-sm) !important; }

/* ════ PROGRESS ════ */
.stProgress > div > div > div > div { background:linear-gradient(90deg,#4f46e5,#7c3aed,#f472b6) !important; border-radius:99px !important; }
.stProgress > div > div { background:rgba(100,108,255,0.1) !important; border-radius:99px !important; }

/* ════ EXPANDER ════ */
[data-testid="stExpander"] { background:transparent !important; border:1px solid var(--border) !important; border-radius:var(--r-sm) !important; }
[data-testid="stExpander"] summary { background:rgba(15,15,34,0.8) !important; color:var(--text-1) !important; }
[data-testid="stExpander"] [data-testid="stVerticalBlock"] { background:rgba(9,9,22,0.6) !important; }

/* ════ CUSTOM BOXES ════ */
.somo-notify {
  background:linear-gradient(135deg,rgba(100,108,255,0.12),rgba(167,139,250,0.08));
  border:1px solid rgba(100,108,255,0.3); border-radius:var(--r);
  padding:13px 18px; color:#c7d2fe; font-weight:600; font-size:14px;
  margin:10px 0; display:flex; align-items:center; gap:10px;
  animation:slideup 0.35s ease;
}
.somo-success {
  background:linear-gradient(135deg,rgba(52,211,153,0.12),rgba(5,150,105,0.08));
  border:1px solid rgba(52,211,153,0.3); color:#6ee7b7;
  border-radius:var(--r); padding:13px 18px; font-weight:600;
  font-size:14px; margin:10px 0; animation:slideup 0.35s ease;
}
@keyframes slideup { from{opacity:0;transform:translateY(-8px)} to{opacity:1;transform:translateY(0)} }

/* ════ API BADGES ════ */
.api-badge {
  display:inline-flex; align-items:center; gap:6px;
  padding:4px 11px; border-radius:99px; font-size:11px;
  font-weight:600; font-family:var(--fmono); letter-spacing:0.5px; border:1px solid;
}
.api-groq    { background:rgba(255,166,0,0.1);   color:#fbbf24; border-color:rgba(251,191,36,0.3); }
.api-gemini  { background:rgba(52,211,153,0.1);  color:#34d399; border-color:rgba(52,211,153,0.3); }
.api-cohere  { background:rgba(56,189,248,0.1);  color:#38bdf8; border-color:rgba(56,189,248,0.3); }
.api-mistral { background:rgba(244,114,182,0.1); color:#f472b6; border-color:rgba(244,114,182,0.3); }
.api-dot {
  width:6px; height:6px; border-radius:50%;
  animation:bdot 2s ease-in-out infinite;
}
@keyframes bdot { 0%,100%{opacity:1} 50%{opacity:0.3} }
.api-groq .api-dot    { background:#fbbf24; }
.api-gemini .api-dot  { background:#34d399; }
.api-cohere .api-dot  { background:#38bdf8; }
.api-mistral .api-dot { background:#f472b6; }

/* ════ TEMPLATES ════ */
.tmpl-card {
  background:var(--bg-card); border:1px solid var(--border);
  border-radius:var(--r); padding:20px; margin-bottom:12px;
  transition:all 0.25s; position:relative; overflow:hidden;
}
.tmpl-card::before { content:''; position:absolute; top:0; left:0; width:3px; height:100%; background:linear-gradient(180deg,var(--accent),#a78bfa); }
.tmpl-card:hover { border-color:var(--border-h); transform:translateX(4px); }
.tmpl-tag { display:inline-block; padding:3px 11px; border-radius:99px; font-size:10px; font-weight:700; letter-spacing:1px; margin-bottom:8px; font-family:var(--fmono); }
.tag-excel  { background:rgba(52,211,153,0.1);  color:#34d399; border:1px solid rgba(52,211,153,0.25); }
.tag-word   { background:rgba(56,189,248,0.1);  color:#38bdf8; border:1px solid rgba(56,189,248,0.25); }
.tag-code   { background:rgba(251,191,36,0.1);  color:#fbbf24; border:1px solid rgba(251,191,36,0.25); }
.tag-html   { background:rgba(244,114,182,0.1); color:#f472b6; border:1px solid rgba(244,114,182,0.25); }
.tmpl-title { font-size:14px; font-weight:700; color:var(--text-1); margin-bottom:4px; font-family:var(--fhead); }
.tmpl-desc  { font-size:11.5px; color:var(--text-3); line-height:1.55; }

/* ════ HISTORY ════ */
.hist-msg { border-left:3px solid; border-radius:0 var(--r) var(--r) 0; padding:11px 15px; margin:7px 0; font-size:13px; }
.hist-user { background:rgba(100,108,255,0.07); border-color:var(--accent); }
.hist-ai   { background:rgba(52,211,153,0.06);  border-color:#34d399; }
.hist-role { font-size:10px; font-weight:700; letter-spacing:1.5px; text-transform:uppercase; margin-bottom:4px; font-family:var(--fmono); }
.hist-user .hist-role { color:#818cf8; }
.hist-ai   .hist-role { color:#34d399; }
.hist-body { color:var(--text-2); line-height:1.55; }
.profile-stat { background:var(--bg-card); border:1px solid var(--border); border-radius:var(--r); padding:20px 16px; text-align:center; }

/* ════ FOOTER ════ */
.somo-footer { text-align:center; padding:44px 20px 22px; border-top:1px solid var(--border); margin-top:60px; }
.f-title { font-size:18px; font-weight:800; color:var(--text-1); margin-bottom:9px; font-family:var(--fhead); }
.f-sub   { font-size:12.5px; color:var(--text-3); margin-bottom:5px; }

/* ════ MOBILE ════ */
@media(max-width:768px) {
  .somo-hero { padding:24px 16px !important; border-radius:18px !important; }
  .somo-hero h1 { font-size:clamp(20px,6vw,30px) !important; }
  [data-testid="stSidebar"] { width:80vw !important; max-width:260px !important; }
}
@media(max-width:480px) {
  .somo-hero h1 { font-size:20px !important; }
  [data-testid="stSidebar"] { width:85vw !important; }
}

/* ════ GLOBAL DARK ════ */
.stApp,[data-testid="stMainBlockContainer"],
[data-testid="stAppViewContainer"],[data-testid="stMain"] {
  background: var(--bg-0) !important;
}
</style>
""", unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════════════════
# CONSTANTS & CONFIG
# ════════════════════════════════════════════════════════════════════════════════

API_CONFIGS: Dict[str, Dict] = {
    "groq": {
        "name": "Groq", "icon": "⚡",
        "model": "llama-3.3-70b-versatile",
        "desc": "Llama 3.3 · Ultra Fast",
        "badge_class": "api-groq"
    },
    "gemini": {
        "name": "Gemini", "icon": "✨",
        "model": "gemini-2.0-flash",
        "desc": "Google Gemini 2.0",
        "badge_class": "api-gemini"
    },
    "cohere": {
        "name": "Cohere", "icon": "🔮",
        "model": "command-r-plus",
        "desc": "Command R+",
        "badge_class": "api-cohere"
    },
    "mistral": {
        "name": "Mistral", "icon": "🌪",
        "model": "mistral-large-latest",
        "desc": "Mistral Large",
        "badge_class": "api-mistral"
    },
}

MIME_TYPES: Dict[str, str] = {
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "py":   "text/x-python",
    "html": "text/html",
    "csv":  "text/csv",
}

PROVIDER_ORDER = ["groq", "cohere", "mistral", "gemini"]  # gemini oxirida — quota muammosi

JSON_RE   = re.compile(r'\{.*\}', re.DOTALL)
CODE_KW   = re.compile(r'(python|kod|script|bot|function|dastur)', re.IGNORECASE)

# ════════════════════════════════════════════════════════════════════════════════
# SECURITY: COOKIES
# ════════════════════════════════════════════════════════════════════════════════

if HAS_COOKIES:
    _cpw = ""
    try:
        _cpw = str(st.secrets.get("COOKIE_PASSWORD", "somo_secure_2026")).strip()
    except Exception:
        _cpw = "somo_secure_2026"
    cookies = EncryptedCookieManager(password=_cpw)
    if not cookies.ready():
        st.stop()
else:
    cookies = {}

# ════════════════════════════════════════════════════════════════════════════════
# SECURITY: PASSWORD HASHING
# ════════════════════════════════════════════════════════════════════════════════

def hash_password(password: str) -> str:
    """Parolni bcrypt bilan hash qilish."""
    if HAS_BCRYPT:
        return bcrypt.hashpw(password.encode(), bcrypt.gensalt(rounds=12)).decode()
    import hashlib
    return hashlib.sha256(password.encode()).hexdigest()


def verify_password(password: str, stored_hash: str) -> bool:
    """Parolni tekshirish — bcrypt va SHA-256 ni qo'llab-quvvatlaydi."""
    if not stored_hash:
        return False
    if HAS_BCRYPT:
        try:
            if stored_hash.startswith(("$2b$", "$2a$")):
                return bcrypt.checkpw(password.encode(), stored_hash.encode())
        except Exception:
            pass
    import hashlib
    return hashlib.sha256(password.encode()).hexdigest() == stored_hash

# ════════════════════════════════════════════════════════════════════════════════
# SECRETS
# ════════════════════════════════════════════════════════════════════════════════

@lru_cache(maxsize=32)
def get_secret(key: str) -> str:
    """API key olish — secrets va env dan."""
    try:
        val = st.secrets.get(key, "")
        if val:
            return str(val).strip()
    except Exception:
        pass
    for section in ["keys", "api", "api_keys", "APIs"]:
        try:
            val = st.secrets[section][key]
            if val:
                return str(val).strip()
        except Exception:
            pass
    return str(os.environ.get(key, "")).strip()

# ════════════════════════════════════════════════════════════════════════════════
# AI CLIENTS
# ════════════════════════════════════════════════════════════════════════════════

@st.cache_resource
def init_ai_clients() -> Dict[str, Any]:
    """Barcha AI clientlarni initialize qilish."""
    clients: Dict[str, Any] = {}
    init_errors: List[str]  = []

    if HAS_GROQ:
        try:
            key = get_secret("GROQ_API_KEY")
            if key:
                clients["groq"] = Groq(api_key=key)
            else:
                init_errors.append("GROQ_API_KEY topilmadi")
        except Exception as e:
            init_errors.append(f"Groq: {e}")

    if HAS_GEMINI:
        try:
            key = get_secret("GEMINI_API_KEY")
            if key:
                genai.configure(api_key=key)
                for model_name in ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-pro"]:
                    try:
                        clients["gemini"] = genai.GenerativeModel(model_name)
                        break
                    except Exception as me:
                        if "not found" in str(me).lower() or "404" in str(me):
                            continue
                        raise
            else:
                init_errors.append("GEMINI_API_KEY topilmadi")
        except Exception as e:
            init_errors.append(f"Gemini: {e}")

    if HAS_COHERE:
        try:
            key = get_secret("COHERE_API_KEY")
            if key:
                # v2 SDK (cohere>=5.0)
                if _COHERE_V2:
                    clients["cohere"] = cohere.ClientV2(api_key=key)
                else:
                    clients["cohere"] = cohere.Client(api_key=key)
            else:
                init_errors.append("COHERE_API_KEY topilmadi")
        except Exception as e:
            init_errors.append(f"Cohere: {e}")

    if HAS_MISTRAL:
        try:
            key = get_secret("MISTRAL_API_KEY")
            if key:
                clients["mistral"] = Mistral(api_key=key)
            else:
                init_errors.append("MISTRAL_API_KEY topilmadi")
        except Exception as e:
            init_errors.append(f"Mistral: {e}")

    # Hech bir client yo'q — xatolarni saqlash (keyinroq ko'rsatiladi)
    if not clients:
        import json as _json
        try:
            with open("/tmp/somo_init_errors.json", "w") as f:
                _json.dump(init_errors, f)
        except Exception:
            pass

    return clients


ai_clients = init_ai_clients()

# Agar hech bir client yo'q bo'lsa — init xatolarini ko'rsatish
if not ai_clients:
    try:
        import json as _j
        _errs = _j.load(open("/tmp/somo_init_errors.json"))
        st.error(
            "⚠️ **Hech bir AI client ishga tushmadi!**\n\n"
            + "\n".join(f"• {e}" for e in _errs)
            + "\n\nStreamlit `secrets.toml` da API kalitlarni tekshiring:\n"
            "```toml\nGROQ_API_KEY = \"gsk_...\"\nGEMINI_API_KEY = \"AIza...\"\n```"
        )
    except Exception:
        st.error("⚠️ Hech bir AI client ishga tushmadi! API kalitlarni tekshiring.")

# ════════════════════════════════════════════════════════════════════════════════
# DATABASE
# ════════════════════════════════════════════════════════════════════════════════

@st.cache_resource
def get_db():
    """Google Sheets ulanish."""
    try:
        if not HAS_OAUTH:
            return None, None, None
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive",
        ]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(
            st.secrets.get("gcp_service_account", {}), scope
        )
        gc = gspread.authorize(creds)
        ss = gc.open("Somo_Users")
        user_ws = ss.sheet1

        try:
            chat_ws = ss.worksheet("ChatHistory")
        except Exception:
            chat_ws = ss.add_worksheet("ChatHistory", 5000, 6)
            chat_ws.append_row(["Timestamp", "Username", "Role", "Message", "Intent", "Provider"])

        try:
            fb_ws = ss.worksheet("Feedback")
        except Exception:
            fb_ws = ss.add_worksheet("Feedback", 1000, 8)
            fb_ws.append_row(["Timestamp", "Username", "Rating", "Category", "Message", "Email", "Status", "Files"])

        return user_ws, chat_ws, fb_ws

    except Exception:
        return None, None, None


user_db, chat_db, fb_db = get_db()


@st.cache_data(ttl=120)
def get_all_users() -> List[Dict]:
    """Foydalanuvchilarni cache bilan olish (120s TTL)."""
    if user_db:
        try:
            return user_db.get_all_records()
        except Exception:
            return []
    return []

# ════════════════════════════════════════════════════════════════════════════════
# RATE-LIMITED LOGGING
# ════════════════════════════════════════════════════════════════════════════════

_last_log: Dict[str, float] = {}


def db_log(username: str, role: str, content: str,
           intent: str = "chat", provider: str = "groq") -> None:
    """Google Sheets ga rate-limited yozuv."""
    global _last_log
    now   = time.time()
    ukey  = f"{username}_{role}"
    if now - _last_log.get(ukey, 0) < 3:
        return
    _last_log[ukey] = now
    if chat_db:
        try:
            chat_db.append_row([
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                username, role, content[:400], intent, provider
            ])
        except Exception:
            pass

# ════════════════════════════════════════════════════════════════════════════════
# DOCUMENT PROCESSING
# ════════════════════════════════════════════════════════════════════════════════

def process_document(file_obj) -> str:
    """PDF yoki DOCX faylni matn sifatida o'qish."""
    try:
        if file_obj.type == "application/pdf" and HAS_PDF:
            reader = PdfReader(file_obj)
            parts = []
            for page in reader.pages[:100]:
                txt = page.extract_text()
                if txt:
                    parts.append(txt)
            return "\n".join(parts)
        if "wordprocessingml" in file_obj.type and HAS_MAMMOTH:
            return mammoth_extract(file_obj).value
    except Exception as e:
        st.warning(f"⚠️ Fayl o'qishda xato: {e}")
    return ""

# ════════════════════════════════════════════════════════════════════════════════
# INTENT DETECTION
# ════════════════════════════════════════════════════════════════════════════════

def detect_intent(text: str) -> str:
    """Foydalanuvchi niyatini aniqlash."""
    t = text.lower()
    if any(k in t for k in ["excel", "xlsx", "jadval", "byudjet", "budget",
                              "spreadsheet", "finance", "moliya", "hisobot",
                              "salary", "ish haqi", "sotish", "savdo", "xarajat",
                              "daromad", "inventar", "formula"]):
        return "excel"
    if any(k in t for k in ["word", "docx", "hujjat", "rezyume", "resume", "cv",
                              "shartnoma", "contract", "ariza", "maktub", "letter",
                              "kurs ishi", "referat", "diplom", "biznes reja"]):
        return "word"
    if any(k in t for k in ["html", "website", "veb sahifa", "landing page"]):
        return "html"
    if any(k in t for k in ["csv", "dataset", "comma separated"]):
        return "csv"
    if CODE_KW.search(t):
        return "code"
    return "chat"

# ════════════════════════════════════════════════════════════════════════════════
# AI CORE FUNCTIONS
# ════════════════════════════════════════════════════════════════════════════════

def call_ai(
    messages: List[Dict],
    temperature: float = 0.6,
    max_tokens: int = 3000,
    provider: str = "groq"
) -> Tuple[str, str]:
    """AI ga so'rov yuborish — fallback bilan."""
    order = [provider] + [p for p in PROVIDER_ORDER if p != provider]

    for prov in order:
        if prov not in ai_clients:
            continue
        try:
            if prov == "groq":
                resp = ai_clients["groq"].chat.completions.create(
                    messages=messages,
                    model=API_CONFIGS["groq"]["model"],
                    temperature=min(temperature, 1.0),
                    max_tokens=max_tokens,
                )
                return resp.choices[0].message.content, "groq"

            elif prov == "gemini":
                sys_msg  = next((m["content"] for m in messages if m["role"] == "system"), "")
                user_msgs = [m for m in messages if m["role"] != "system"]
                hist     = [
                    {"role": "user" if m["role"] == "user" else "model",
                     "parts": [m["content"]]}
                    for m in user_msgs[:-1]
                ]
                last = user_msgs[-1]["content"] if user_msgs else ""
                if sys_msg:
                    last = f"[System: {sys_msg}]\n\n{last}"
                chat = ai_clients["gemini"].start_chat(history=hist)
                return chat.send_message(last).text, "gemini"

            elif prov == "cohere":
                # Cohere v2 SDK (cohere>=5.0): ClientV2 + messages=[] format
                try:
                    sys_msg   = next((m["content"] for m in messages if m["role"] == "system"), "")
                    user_msgs = [m for m in messages if m["role"] != "system"]
                    co_client = ai_clients["cohere"]

                    msgs_v2 = []
                    if sys_msg:
                        msgs_v2.append({"role": "system", "content": sys_msg})
                    for m in user_msgs[:-1]:
                        msgs_v2.append({"role": m["role"], "content": m["content"]})
                    if user_msgs:
                        msgs_v2.append({"role": "user", "content": user_msgs[-1]["content"]})

                    resp = co_client.chat(
                        model=API_CONFIGS["cohere"]["model"],
                        messages=msgs_v2,
                    )
                    # v2 javob: resp.message.content[0].text
                    if hasattr(resp, "message") and resp.message:
                        content_list = getattr(resp.message, "content", [])
                        if content_list:
                            txt = content_list[0].text if hasattr(content_list[0], "text") else str(content_list[0])
                            if txt:
                                return txt, "cohere"
                    # v1 fallback
                    if hasattr(resp, "text") and resp.text:
                        return resp.text, "cohere"
                except Exception:
                    pass
                continue

            elif prov == "mistral":
                resp = ai_clients["mistral"].chat.complete(
                    model=API_CONFIGS["mistral"]["model"],
                    messages=messages,
                    temperature=min(temperature, 1.0),
                    max_tokens=max_tokens,
                )
                return resp.choices[0].message.content, "mistral"

        except Exception:
            continue

    return "❌ Hech bir AI xizmati javob bermadi.", "none"


def call_ai_stream(
    messages: List[Dict],
    temperature: float = 0.6,
    max_tokens: int = 3000,
    provider: str = "groq"
):
    """Streaming AI so'rovi — har bir provider uchun to'liq fallback va xato log."""
    order  = [provider] + [p for p in PROVIDER_ORDER if p != provider]
    errors: Dict[str, str] = {}

    for prov in order:
        if prov not in ai_clients:
            errors[prov] = "client yo'q"
            continue
        try:
            if prov == "groq":
                stream = ai_clients["groq"].chat.completions.create(
                    messages=messages,
                    model=API_CONFIGS["groq"]["model"],
                    temperature=min(temperature, 1.0),
                    max_tokens=max_tokens,
                    stream=True,
                )
                yielded = False
                for chunk in stream:
                    delta = chunk.choices[0].delta.content
                    if delta:
                        yielded = True
                        yield delta, "groq"
                if yielded:
                    return
                # Bo'sh javob — fallback
                errors["groq"] = "bo'sh javob"
                continue

            elif prov == "gemini":
                sys_msg   = next((m["content"] for m in messages if m["role"] == "system"), "")
                user_msgs = [m for m in messages if m["role"] != "system"]
                hist      = [
                    {"role": "user" if m["role"] == "user" else "model",
                     "parts": [m["content"]]}
                    for m in user_msgs[:-1]
                ]
                last = user_msgs[-1]["content"] if user_msgs else ""
                if sys_msg:
                    last = f"[System: {sys_msg}]\n\n{last}"
                chat_obj = ai_clients["gemini"].start_chat(history=hist)
                yielded  = False
                for chunk in chat_obj.send_message(last, stream=True):
                    if chunk.text:
                        yielded = True
                        yield chunk.text, "gemini"
                if yielded:
                    return
                errors["gemini"] = "bo'sh javob"
                continue

            elif prov == "cohere":
                # Cohere v2 SDK — stream=True parametri bilan
                sys_msg   = next((m["content"] for m in messages if m["role"] == "system"), "")
                user_msgs = [m for m in messages if m["role"] != "system"]
                co_client = ai_clients["cohere"]

                msgs_v2 = []
                if sys_msg:
                    msgs_v2.append({"role": "system", "content": sys_msg})
                for m in user_msgs[:-1]:
                    msgs_v2.append({"role": m["role"], "content": m["content"]})
                if user_msgs:
                    msgs_v2.append({"role": "user", "content": user_msgs[-1]["content"]})

                # 1. Cohere v2 streaming (stream=True parametri bilan)
                yielded = False
                try:
                    for event in co_client.chat(
                        model=API_CONFIGS["cohere"]["model"],
                        messages=msgs_v2,
                        stream=True,
                    ):
                        # v2 stream event turlari
                        evt_type = getattr(event, "type", None)
                        if evt_type == "content-delta":
                            # event.delta.message.content.text
                            try:
                                txt = event.delta.message.content.text
                                if txt:
                                    yielded = True
                                    yield txt, "cohere"
                                    continue
                            except Exception:
                                pass
                        # fallback: to'g'ridan-to'g'ri .text atributi
                        raw_txt = getattr(event, "text", None)
                        if raw_txt:
                            yielded = True
                            yield raw_txt, "cohere"
                    if yielded:
                        return
                except Exception:
                    pass

                # 2. Non-stream fallback — butun javob birdan
                try:
                    resp = co_client.chat(
                        model=API_CONFIGS["cohere"]["model"],
                        messages=msgs_v2,
                    )
                    result = None
                    if hasattr(resp, "message") and resp.message:
                        content_list = getattr(resp.message, "content", [])
                        if content_list:
                            result = content_list[0].text if hasattr(content_list[0], "text") else str(content_list[0])
                    if not result and hasattr(resp, "text"):
                        result = resp.text
                    if result:
                        yield result, "cohere"
                        return
                except Exception as e2:
                    errors["cohere"] = str(e2)[:100]
                continue

            elif prov == "mistral":
                yielded = False
                for chunk in ai_clients["mistral"].chat.stream(
                    model=API_CONFIGS["mistral"]["model"],
                    messages=messages,
                    temperature=min(temperature, 1.0),
                    max_tokens=max_tokens,
                ):
                    delta = chunk.data.choices[0].delta.content
                    if delta:
                        yielded = True
                        yield delta, "mistral"
                if yielded:
                    return
                errors["mistral"] = "bo'sh javob"
                continue

        except Exception as e:
            errors[prov] = str(e)[:120]
            continue

    # Barcha providerlar muvaffaqiyatsiz — non-stream fallback sinab ko'rish
    fallback_resp, fallback_prov = call_ai(messages, temperature, max_tokens, provider)
    if fallback_prov != "none":
        yield fallback_resp, fallback_prov
        return

    # Hech narsa ishlamadi — xato detallarini ko'rsatish
    err_lines = "\n".join([f"• **{p}**: {e}" for p, e in errors.items()])
    yield (
        f"❌ **Hech bir AI javob bermadi.**\n\n"
        f"**Xatolar:**\n{err_lines}\n\n"
        f"💡 *Tekshiring: API kalitlar to'g'ri sozlanganmi? "
        f"Rate limit bo'lishi mumkin — bir daqiqa kuting.*"
    ), "none"

# ════════════════════════════════════════════════════════════════════════════════
# FILE GENERATORS
# ════════════════════════════════════════════════════════════════════════════════

def gen_excel(prompt: str, temp: float = 0.15,
              provider: str = "groq") -> Tuple[Optional[bytes], str]:
    """Excel fayl yaratish."""
    if not HAS_OPENPYXL:
        return None, "openpyxl o'rnatilmagan"

    sys_p = """Sen Excel fayl strukturasi uchun JSON qaytaruvchi ekspertsan.
FAQAT quyidagi JSON formatini qaytar:
{
  "title": "Fayl nomi",
  "sheets": [
    {
      "name": "Varaq nomi",
      "headers": ["Sarlavha1", "Sarlavha2"],
      "header_color": "4F46E5",
      "rows": [["qiymat1", "qiymat2"]],
      "col_widths": [25, 15],
      "row_height": 20
    }
  ]
}
Kamida 15-20 satr. Excel formulalardan foydalaning: SUM, AVERAGE, IF. FAQAT JSON."""

    raw, _ = call_ai(
        [{"role": "system", "content": sys_p},
         {"role": "user",   "content": prompt}],
        temperature=temp, max_tokens=4000, provider=provider
    )
    raw = re.sub(r'```json|```', '', raw).strip()
    m   = JSON_RE.search(raw)
    if not m:
        return None, "JSON topilmadi"
    try:
        data = json.loads(m.group())
    except Exception as e:
        return None, f"JSON xatosi: {e}"

    PALETTES = [
        ("4F46E5", "EEF2FF"), ("059669", "ECFDF5"),
        ("D97706", "FFFBEB"), ("DC2626", "FEF2F2"),
        ("0891B2", "ECFEFF"), ("7C3AED", "F5F3FF"),
    ]

    wb = Workbook()
    wb.remove(wb.active)

    for si, sh in enumerate(data.get("sheets", [])):
        ws       = wb.create_sheet(title=sh.get("name", "Sheet")[:31])
        headers  = sh.get("headers", [])
        hcolor   = sh.get("header_color", PALETTES[si % len(PALETTES)][0])
        _, rcolor = PALETTES[si % len(PALETTES)]
        rows     = sh.get("rows", [])
        widths   = sh.get("col_widths", [])
        rh       = sh.get("row_height", 20)

        if headers:
            end_col = max(len(headers), 1)
            ws.merge_cells(start_row=1, start_column=1,
                           end_row=1, end_column=end_col)
            tc = ws.cell(row=1, column=1, value=sh.get("name", "Hisobot"))
            tc.font      = Font(name="Calibri", bold=True, size=13, color="FFFFFF")
            tc.fill      = PatternFill("solid", fgColor=hcolor)
            tc.alignment = Alignment(horizontal="center", vertical="center")
            ws.row_dimensions[1].height = 30

            th = Side(style="medium", color="FFFFFF")
            for ci, h in enumerate(headers, 1):
                c = ws.cell(row=2, column=ci, value=h)
                c.font      = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
                c.fill      = PatternFill("solid", fgColor=hcolor)
                c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                c.border    = Border(left=th, right=th, top=th, bottom=th)
            ws.row_dimensions[2].height = 24

        td = Side(style="thin", color="D1D5DB")
        for ri, row in enumerate(rows, 3):
            bg = "FFFFFF" if ri % 2 else rcolor
            for ci, val in enumerate(row, 1):
                c = ws.cell(row=ri, column=ci)
                if isinstance(val, str) and val.startswith("="):
                    c.value = val
                else:
                    try:
                        if isinstance(val, str) and re.match(r'^-?\d+(\.\d+)?$', val.strip()):
                            v = float(val)
                            c.value = int(v) if v == int(v) else v
                        else:
                            c.value = val
                    except Exception:
                        c.value = val
                c.fill      = PatternFill("solid", fgColor=bg)
                c.border    = Border(left=td, right=td, top=td, bottom=td)
                c.font      = Font(name="Calibri", size=10)
                c.alignment = Alignment(vertical="center", wrap_text=True)
            ws.row_dimensions[ri].height = rh

        for ci, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(ci)].width = max(int(w), 8)
        if not widths and headers:
            for ci in range(1, len(headers) + 1):
                ws.column_dimensions[get_column_letter(ci)].width = 18
        ws.freeze_panes = "A3"

    if not wb.sheetnames:
        wb.create_sheet("Ma'lumotlar")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    safe = re.sub(r'[^\w\s-]', '', data.get("title", "somo")).strip().replace(' ', '_')
    return buf.getvalue(), f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"


def gen_word(prompt: str, temp: float = 0.4,
             provider: str = "mistral") -> Tuple[Optional[bytes], str]:
    """Word hujjat yaratish."""
    if not HAS_DOCX:
        return None, "python-docx o'rnatilmagan"

    sys_p = """Sen professional Word hujjat strukturasi JSON qaytaruvchi ekspertsan.
FAQAT JSON:
{
  "title": "Sarlavha",
  "sections": [
    {"type": "heading1",  "text": "..."},
    {"type": "heading2",  "text": "..."},
    {"type": "paragraph", "text": "..."},
    {"type": "bullet",    "items": ["1", "2"]},
    {"type": "numbered",  "items": ["A", "B"]},
    {"type": "table", "headers": ["H1","H2"], "rows": [["v1","v2"]]}
  ]
}
Kamida 10-14 bo'lim. Faqat JSON."""

    raw, _ = call_ai(
        [{"role": "system", "content": sys_p},
         {"role": "user",   "content": prompt}],
        temperature=temp, max_tokens=4000, provider=provider
    )
    raw = re.sub(r'```json|```', '', raw).strip()
    m   = JSON_RE.search(raw)
    if not m:
        return None, "Struktura topilmadi"
    try:
        data = json.loads(m.group())
    except Exception as e:
        return None, f"JSON: {e}"

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2.5)

    # Title
    tp  = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(data.get("title", "Hujjat"))
    run.bold           = True
    run.font.size      = Pt(20)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    run.font.name      = "Calibri"

    # Date line
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"{datetime.now().strftime('%d.%m.%Y')} — Somo AI")
    dr.font.size      = Pt(10)
    dr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    dr.font.name      = "Calibri"

    for sec in data.get("sections", []):
        t = sec.get("type", "paragraph")

        if t == "heading1":
            h = doc.add_heading(sec.get("text", ""), level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                h.runs[0].font.name      = "Calibri"

        elif t == "heading2":
            h = doc.add_heading(sec.get("text", ""), level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
                h.runs[0].font.name      = "Calibri"

        elif t == "paragraph":
            p  = doc.add_paragraph()
            r  = p.add_run(sec.get("text", ""))
            r.font.size = Pt(11)
            r.font.name = "Calibri"
            p.paragraph_format.space_after    = Pt(8)
            p.paragraph_format.line_spacing   = Pt(16)

        elif t == "bullet":
            for item in sec.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(item)
                r.font.size = Pt(11)
                r.font.name = "Calibri"

        elif t == "numbered":
            for item in sec.get("items", []):
                p = doc.add_paragraph(style="List Number")
                r = p.add_run(item)
                r.font.size = Pt(11)
                r.font.name = "Calibri"

        elif t == "table":
            hdrs = sec.get("headers", [])
            rws  = sec.get("rows", [])
            if hdrs:
                tbl = doc.add_table(rows=1 + len(rws), cols=len(hdrs))
                tbl.style = "Table Grid"
                hrow = tbl.rows[0]
                for ci, h in enumerate(hdrs):
                    cell = hrow.cells[ci]
                    cell.text = h
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold      = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        cell.paragraphs[0].runs[0].font.name      = "Calibri"
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd  = OxmlElement("w:shd")
                    shd.set(qn("w:val"),   "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"),  "4F46E5")
                    tcPr.append(shd)
                for ri, rdata in enumerate(rws):
                    rcells = tbl.rows[ri + 1].cells
                    for ci, val in enumerate(rdata):
                        if ci < len(rcells):
                            rcells[ci].text = str(val)
                            if rcells[ci].paragraphs[0].runs:
                                rcells[ci].paragraphs[0].runs[0].font.size = Pt(10)
                                rcells[ci].paragraphs[0].runs[0].font.name = "Calibri"
                doc.add_paragraph()

    # Footer
    fp  = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr  = fp.add_run(f"© {datetime.now().year} Somo AI  |  {data.get('title','')}")
    fr.font.size      = Pt(9)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    fr.font.name      = "Calibri"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe = re.sub(r'[^\w\s-]', '', data.get("title", "somo")).strip().replace(' ', '_')
    return buf.getvalue(), f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"


def gen_code(prompt: str, temp: float = 0.12,
             provider: str = "cohere") -> Tuple[bytes, str]:
    """Python kod yaratish."""
    sys_p = ("Sen tajribali Python dasturchi. Professional, to'liq ishlaydigan kod yoz. "
             "Error handling, docstring, logging bilan. FAQAT Python kodi — markdown yo'q.")
    raw, _ = call_ai(
        [{"role": "system", "content": sys_p},
         {"role": "user",   "content": prompt}],
        temperature=temp, max_tokens=3500, provider=provider
    )
    raw  = re.sub(r'```python|```py|```', '', raw).strip()
    safe = re.sub(r'[^\w]', '_', prompt[:30]).strip('_')
    return raw.encode('utf-8'), f"{safe}_{datetime.now().strftime('%H%M%S')}.py"


def gen_html(prompt: str, temp: float = 0.5,
             provider: str = "gemini") -> Tuple[bytes, str]:
    """HTML sahifa yaratish."""
    sys_p = ("Sen professional frontend developer. Chiroyli, zamonaviy HTML/CSS/JS sahifa yarat. "
             "Dark theme, glassmorphism, animatsiyalar. FAQAT HTML kodi — markdown yo'q.")
    raw, _ = call_ai(
        [{"role": "system", "content": sys_p},
         {"role": "user",   "content": prompt}],
        temperature=temp, max_tokens=4000, provider=provider
    )
    raw  = re.sub(r'```html|```', '', raw).strip()
    safe = re.sub(r'[^\w]', '_', prompt[:25]).strip('_')
    return raw.encode('utf-8'), f"{safe}_{datetime.now().strftime('%H%M%S')}.html"


def gen_csv(prompt: str, temp: float = 0.3,
            provider: str = "mistral") -> Tuple[bytes, str]:
    """CSV dataset yaratish."""
    sys_p = ("Sen ma'lumotlar mutaxassisi. FAQAT CSV format. "
             "Birinchi satr sarlavha. Kamida 25 satr. Hech qanday tushuntirma yo'q.")
    raw, _ = call_ai(
        [{"role": "system", "content": sys_p},
         {"role": "user",   "content": prompt}],
        temperature=temp, max_tokens=3000, provider=provider
    )
    raw  = re.sub(r'```csv|```', '', raw).strip()
    safe = re.sub(r'[^\w]', '_', prompt[:25]).strip('_')
    return raw.encode('utf-8'), f"{safe}_{datetime.now().strftime('%H%M%S')}.csv"

# ════════════════════════════════════════════════════════════════════════════════
# UI HELPERS
# ════════════════════════════════════════════════════════════════════════════════

def api_badge(provider: str) -> str:
    """API holat badge HTML."""
    cfg = API_CONFIGS.get(provider, API_CONFIGS["groq"])
    return (f'<span class="api-badge {cfg["badge_class"]}">'
            f'<span class="api-dot"></span>{cfg["icon"]} {cfg["name"]}</span>')


def download_btn(file_bytes: bytes, filename: str, label: str) -> None:
    """Download tugmasi."""
    ext  = filename.rsplit('.', 1)[-1]
    mime = MIME_TYPES.get(ext, "application/octet-stream")
    st.markdown(f'<div class="somo-success">✅ {label} fayl tayyor!</div>',
                unsafe_allow_html=True)
    st.download_button(
        f"⬇️  {filename}", file_bytes, filename, mime,
        use_container_width=True, type="primary",
        key=f"dl_{filename}_{time.time()}"
    )


def provider_selector(session_key: str, widget_key: str,
                      label: str = "🤖 AI") -> str:
    """Provider tanlash selectbox — xavfsiz."""
    avail  = [p for p in PROVIDER_ORDER if p in ai_clients]
    if not avail:
        st.warning("⚠️ Hech bir AI provider aktiv emas!")
        return "groq"
    labels = [f"{API_CONFIGS[p]['icon']} {API_CONFIGS[p]['name']}" for p in avail]
    curr   = st.session_state.get(session_key, avail[0])
    if curr not in avail:
        curr = avail[0]
    idx    = avail.index(curr)
    sel    = st.selectbox(label, labels, index=idx, key=widget_key)
    chosen = avail[labels.index(sel)]
    st.session_state[session_key] = chosen
    return chosen

# ════════════════════════════════════════════════════════════════════════════════
# SESSION STATE INIT
# ════════════════════════════════════════════════════════════════════════════════

_SESSION_DEFAULTS: Dict[str, Any] = {
    "logged_in":       False,
    "username":        "",
    "login_time":      datetime.now(),
    "page":            "home",
    "messages":        [],
    "uploaded_text":   "",
    "temperature":     0.6,
    "files_count":     0,
    "ai_style":        "Aqlli yordamchi",
    "chat_provider":   "groq",
    "excel_provider":  "groq",
    "word_provider":   "mistral",
    "code_provider":   "cohere",
    "html_provider":   "gemini",
    "csv_provider":    "mistral",
    "analyze_provider":"gemini",
}

for _k, _v in _SESSION_DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ════════════════════════════════════════════════════════════════════════════════
# SESSION RESTORE (cookie)
# ════════════════════════════════════════════════════════════════════════════════

if not st.session_state.logged_in:
    _saved = cookies.get("somo_user_session") if HAS_COOKIES else None
    if _saved and user_db:
        try:
            _recs = get_all_users()
            _ud   = next((r for r in _recs if str(r.get("username")) == _saved), None)
            if _ud and str(_ud.get("status", "")).lower() == "active":
                st.session_state.logged_in  = True
                st.session_state.username   = _saved
                st.session_state.login_time = datetime.now()
        except Exception:
            pass


def logout() -> None:
    """Tizimdan chiqish."""
    if HAS_COOKIES:
        try:
            cookies["somo_user_session"] = ""
            cookies.save()
        except Exception:
            pass
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.session_state.logged_in = False
    st.rerun()

# ════════════════════════════════════════════════════════════════════════════════
# LOGIN / REGISTER PAGE
# ════════════════════════════════════════════════════════════════════════════════

if not st.session_state.logged_in:
    st.markdown("""
    <div class="somo-hero" style="text-align:center;padding:72px 40px;">
        <div class="grid-dots"></div>
        <div class="somo-hero-content">
            <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#646cff;
                      margin-bottom:14px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
                ✦ Next Generation AI Platform
            </p>
            <h1 style="font-size:clamp(38px,5.5vw,68px);letter-spacing:-2.5px;">
                🌌 Somo AI <span class="g-text">Ultra Pro Max</span>
            </h1>
            <p style="font-size:16px;color:rgba(255,255,255,0.55);
                      max-width:560px;margin:18px auto 30px;line-height:1.7;">
                Excel · Word · Kod · HTML · CSV — To'rt xil AI bilan har qanday faylni yarating
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    cl, cm, cr = st.columns([1, 2, 1])
    with cm:
        t1, t2, t3 = st.tabs(["🔒  Kirish", "✍️  Ro'yxat", "ℹ️  Haqida"])

        # ── LOGIN ──
        with t1:
            st.markdown('<div style="height:12px"></div>', unsafe_allow_html=True)
            with st.form("login_form", clear_on_submit=False):
                st.markdown('<p class="section-label">Hisobingizga kiring</p>',
                            unsafe_allow_html=True)
                u_inp = st.text_input("Username", placeholder="Username kiriting")
                p_inp = st.text_input("Parol", type="password", placeholder="Parolni kiriting")
                rc, bc = st.columns([1, 2])
                with rc:
                    remember = st.checkbox("Eslab qolish", value=True)
                with bc:
                    do_login = st.form_submit_button("🚀  Kirish",
                                                     use_container_width=True,
                                                     type="primary")
                if do_login:
                    if not u_inp or not p_inp:
                        st.error("❌ Username va parolni kiriting!")
                    elif user_db:
                        try:
                            recs = get_all_users()
                            user = next((r for r in recs
                                         if str(r.get("username")) == u_inp), None)
                            if user and verify_password(p_inp, str(user.get("password", ""))):
                                if str(user.get("status", "")).lower() == "blocked":
                                    st.error("🚫 Hisob bloklangan!")
                                else:
                                    st.session_state.logged_in  = True
                                    st.session_state.username   = u_inp
                                    st.session_state.login_time = datetime.now()
                                    if remember and HAS_COOKIES:
                                        cookies["somo_user_session"] = u_inp
                                        cookies.save()
                                    st.success("✅ Muvaffaqiyatli!")
                                    time.sleep(0.4)
                                    st.rerun()
                            else:
                                st.error("❌ Login yoki parol noto'g'ri!")
                        except Exception as e:
                            st.error(f"❌ {e}")
                    else:
                        st.error("❌ Baza ulanmagan!")

        # ── REGISTER ──
        with t2:
            st.markdown('<div style="height:12px"></div>', unsafe_allow_html=True)
            with st.form("register_form", clear_on_submit=False):
                st.markdown('<p class="section-label">Yangi hisob yaratish</p>',
                            unsafe_allow_html=True)
                nu = st.text_input("Username", placeholder="Kamida 3 ta belgi")
                np = st.text_input("Parol",    type="password", placeholder="Kamida 6 ta belgi")
                nc = st.text_input("Tasdiqlash",type="password",placeholder="Qayta kiriting")
                agree   = st.checkbox("Foydalanish shartlariga roziman ✅")
                do_reg  = st.form_submit_button("✨  Hisob yaratish",
                                                use_container_width=True, type="primary")
                if do_reg:
                    if not agree:         st.error("❌ Shartlarga rozilik bering!")
                    elif len(nu) < 3:     st.error("❌ Username kamida 3 belgi!")
                    elif len(np) < 6:     st.error("❌ Parol kamida 6 belgi!")
                    elif np != nc:        st.error("❌ Parollar mos emas!")
                    elif user_db:
                        try:
                            recs = get_all_users()
                            if any(r.get("username") == nu for r in recs):
                                st.error("❌ Bu username band!")
                            else:
                                user_db.append_row([
                                    nu, hash_password(np),
                                    "active", str(datetime.now()), 0
                                ])
                                get_all_users.clear()
                                st.balloons()
                                st.success("🎉 Muvaffaqiyatli! «Kirish» bo'limiga o'ting.")
                        except Exception as e:
                            st.error(f"❌ {e}")
                    else:
                        st.error("❌ Baza ulanmagan!")

        # ── INFO ──
        with t3:
            st.markdown("""
            <div style="padding:8px 0">
            <p class="section-label">Platformalar</p>
            <p class="section-title" style="font-size:18px;">4 AI · 5 Format · ∞ Imkoniyat</p>
            <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-top:16px;">
                <div style="background:rgba(251,191,36,0.06);border:1px solid rgba(251,191,36,0.2);border-radius:12px;padding:14px;">
                    <p style="color:#fbbf24;font-weight:700;font-size:12px;font-family:'JetBrains Mono',monospace;">⚡ GROQ</p>
                    <p style="color:#50506a;font-size:11px;margin-top:4px;line-height:1.6;">Chat, Excel<br>Llama 3.3 70B · Ultra Fast</p>
                </div>
                <div style="background:rgba(52,211,153,0.06);border:1px solid rgba(52,211,153,0.2);border-radius:12px;padding:14px;">
                    <p style="color:#34d399;font-weight:700;font-size:12px;font-family:'JetBrains Mono',monospace;">✨ GEMINI</p>
                    <p style="color:#50506a;font-size:11px;margin-top:4px;line-height:1.6;">HTML, Tahlil<br>Gemini 2.0 Flash</p>
                </div>
                <div style="background:rgba(56,189,248,0.06);border:1px solid rgba(56,189,248,0.2);border-radius:12px;padding:14px;">
                    <p style="color:#38bdf8;font-weight:700;font-size:12px;font-family:'JetBrains Mono',monospace;">🔮 COHERE</p>
                    <p style="color:#50506a;font-size:11px;margin-top:4px;line-height:1.6;">Kod<br>Command R+</p>
                </div>
                <div style="background:rgba(244,114,182,0.06);border:1px solid rgba(244,114,182,0.2);border-radius:12px;padding:14px;">
                    <p style="color:#f472b6;font-weight:700;font-size:12px;font-family:'JetBrains Mono',monospace;">🌪 MISTRAL</p>
                    <p style="color:#50506a;font-size:11px;margin-top:4px;line-height:1.6;">Word, CSV<br>Mistral Large</p>
                </div>
            </div>
            <p style="color:#334155;font-size:11px;margin-top:16px;font-family:'JetBrains Mono',monospace;">
                👨‍💻 Usmonov Sodiq · v4.0 · 2026
            </p>
            </div>
            """, unsafe_allow_html=True)

    st.stop()

# ════════════════════════════════════════════════════════════════════════════════
# SIDEBAR  (logged-in)
# ════════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    uname = st.session_state.username
    avail = [p for p in PROVIDER_ORDER if p in ai_clients]

    # ── User card ──
    st.markdown(f"""
    <div style="padding:22px 16px 18px;border-bottom:1px solid rgba(100,108,255,0.12);margin-bottom:6px;">
        <div style="display:flex;align-items:center;gap:12px;margin-bottom:14px;">
            <div style="background:linear-gradient(135deg,#4f46e5,#7c3aed,#f472b6);
                       width:44px;height:44px;border-radius:14px;flex-shrink:0;
                       display:flex;align-items:center;justify-content:center;
                       font-size:18px;font-weight:900;color:white;
                       box-shadow:0 0 20px rgba(100,108,255,0.4);">
                {uname[0].upper()}
            </div>
            <div>
                <div style="font-size:14px;font-weight:700;color:#f0f0ff;font-family:'Syne',sans-serif;">{uname}</div>
                <div style="font-size:10px;color:#34d399;font-weight:600;font-family:'JetBrains Mono',monospace;
                            display:flex;align-items:center;gap:4px;margin-top:2px;">
                    <span style="background:#34d399;width:5px;height:5px;border-radius:50%;display:inline-block;"></span>
                    ONLINE
                </div>
            </div>
        </div>
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;">
            <div style="background:rgba(100,108,255,0.07);border:1px solid rgba(100,108,255,0.12);
                       border-radius:10px;padding:10px;text-align:center;">
                <div style="font-size:17px;font-weight:900;color:#f0f0ff;">{len(st.session_state.messages)}</div>
                <div style="font-size:9px;color:#50506a;text-transform:uppercase;letter-spacing:1.5px;font-family:'JetBrains Mono',monospace;">Xabar</div>
            </div>
            <div style="background:rgba(52,211,153,0.07);border:1px solid rgba(52,211,153,0.12);
                       border-radius:10px;padding:10px;text-align:center;">
                <div style="font-size:17px;font-weight:900;color:#f0f0ff;">{st.session_state.files_count}</div>
                <div style="font-size:9px;color:#50506a;text-transform:uppercase;letter-spacing:1.5px;font-family:'JetBrains Mono',monospace;">Fayl</div>
            </div>
        </div>
        <div style="display:flex;gap:6px;flex-wrap:wrap;margin-top:10px;">
            {''.join([f'<span class="api-badge {API_CONFIGS[p]["badge_class"]}" style="font-size:9px;padding:3px 8px;"><span class="api-dot"></span>{API_CONFIGS[p]["icon"]}</span>' for p in avail])}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Navigation ──
    st.markdown('<p class="section-label" style="padding:8px 14px 4px;font-size:9px;">Navigatsiya</p>',
                unsafe_allow_html=True)

    NAV = [
        ("home",      "🏠",  "Bosh sahifa"),
        ("chat",      "💬",  "Chat AI"),
        ("excel",     "📊",  "Excel Generator"),
        ("word",      "📝",  "Word Generator"),
        ("code",      "💻",  "Kod Generator"),
        ("html",      "🌐",  "HTML Generator"),
        ("csv",       "📋",  "CSV Generator"),
        ("templates", "🎨",  "Shablonlar"),
        ("analyze",   "🔍",  "Hujjat Tahlili"),
        ("history",   "📜",  "Chat Tarixi"),
        ("feedback",  "💌",  "Fikr Bildirish"),
        ("profile",   "👤",  "Profil"),
    ]

    for pid, icon, label in NAV:
        active = st.session_state.page == pid
        if st.button(f"{icon}  {label}", key=f"nav_{pid}",
                     use_container_width=True,
                     type="primary" if active else "secondary"):
            st.session_state.page = pid
            st.rerun()

    st.markdown('<hr class="somo-divider" style="margin:10px 0;">',
                unsafe_allow_html=True)

    # ── Chat sidebar controls ──
    if st.session_state.page == "chat":
        st.markdown('<p class="section-label" style="padding:0 14px 6px;font-size:9px;">Chat Sozlamalari</p>',
                    unsafe_allow_html=True)
        provider_selector("chat_provider", "sb_chat_prov", "🤖 AI Provider")
        st.session_state.temperature = st.slider(
            "🌡  Ijodkorlik", 0.0, 1.0,
            st.session_state.temperature, 0.05, key="sb_temp"
        )
        st.session_state.ai_style = st.selectbox(
            "💬  Uslub",
            ["Aqlli yordamchi", "Do'stona", "Rasmiy ekspert", "Ijodkor", "Texnik"],
            key="sb_style"
        )
        if st.button("🗑  Chatni tozalash", use_container_width=True, key="sb_clr"):
            st.session_state.messages = []
            st.rerun()

        if st.session_state.messages:
            safe_msgs = [{"role": m.get("role", ""), "content": m.get("content", ""),
                          "provider": m.get("provider", "")}
                         for m in st.session_state.messages]
            st.download_button(
                "📥  JSON Export",
                json.dumps(safe_msgs, ensure_ascii=False, indent=2).encode(),
                f"chat_{datetime.now():%Y%m%d}.json",
                use_container_width=True
            )

    # ── API Diagnostics ──
    st.markdown('<hr class="somo-divider" style="margin:8px 0;">',
                unsafe_allow_html=True)
    st.markdown('<p class="section-label" style="padding:0 14px 4px;font-size:9px;">API Holati</p>',
                unsafe_allow_html=True)

    _all_provs = {"groq": "⚡", "cohere": "🔮", "mistral": "🌪", "gemini": "✨"}
    for _p, _ico in _all_provs.items():
        _ok = _p in ai_clients
        _dot_color = "#34d399" if _ok else "#ef4444"
        _status    = "TAYYOR" if _ok else "YO'Q"
        st.markdown(f"""
        <div style="display:flex;align-items:center;justify-content:space-between;
                    padding:5px 14px;font-size:11px;font-family:'JetBrains Mono',monospace;">
            <span style="color:#a0a0c0;">{_ico} {_p.capitalize()}</span>
            <span style="color:{_dot_color};font-weight:700;font-size:10px;">{_status}</span>
        </div>
        """, unsafe_allow_html=True)

    if not ai_clients:
        st.markdown("""
        <div style="margin:8px 14px;padding:10px;background:rgba(239,68,68,0.1);
                    border:1px solid rgba(239,68,68,0.3);border-radius:8px;
                    font-size:11px;color:#fca5a5;font-family:'JetBrains Mono',monospace;">
            ⚠️ Hech bir API kalit<br>topilmadi! secrets.toml<br>ni tekshiring.
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<br>', unsafe_allow_html=True)
    if st.button("🚪  Tizimdan chiqish", use_container_width=True,
                 type="primary", key="sb_logout"):
        logout()

    st.markdown(f"""
    <div style="padding:14px 14px 8px;border-top:1px solid rgba(100,108,255,0.08);
                margin-top:6px;text-align:center;">
        <p style="font-size:9px;color:#2a2a40;line-height:1.7;font-family:'JetBrains Mono',monospace;">
            🌌 SOMO AI · v4.0<br>⚡ GROQ · ✨ GEMINI · 🔮 COHERE · 🌪 MISTRAL<br>© 2026
        </p>
    </div>
    """, unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: HOME
# ════════════════════════════════════════════════════════════════════════════════

if st.session_state.page == "home":
    avail = [p for p in PROVIDER_ORDER if p in ai_clients]
    mins  = (datetime.now() - st.session_state.login_time).seconds // 60

    st.markdown(f"""
    <div class="somo-hero">
        <div class="grid-dots"></div>
        <div class="somo-hero-content">
            <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#646cff;
                      margin-bottom:12px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
                ✦ Somo AI Ultra Pro Max v4.0
            </p>
            <h1>Salom, <span class="g-text">{uname}</span>! 👋</h1>
            <p class="subtitle">Bugun nima yaratmoqchisiz? To'rt xil AI bilan — Excel, Word, Kod, HTML, CSV — hammasini bir joyda.</p>
            <div class="hero-badges">
                {''.join([f'<span class="hero-badge">{API_CONFIGS[p]["icon"]} {API_CONFIGS[p]["name"]}</span>' for p in avail])}
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="stat-row">
        <div class="stat-box"><div class="stat-icon">💬</div><div class="stat-val">{len(st.session_state.messages)}</div><div class="stat-lbl">Xabarlar</div></div>
        <div class="stat-box"><div class="stat-icon">📁</div><div class="stat-val">{st.session_state.files_count}</div><div class="stat-lbl">Fayllar</div></div>
        <div class="stat-box"><div class="stat-icon">⏱</div><div class="stat-val">{mins}</div><div class="stat-lbl">Daqiqa</div></div>
        <div class="stat-box"><div class="stat-icon">🤖</div><div class="stat-val">{len(avail)}</div><div class="stat-lbl">AI Aktiv</div></div>
        <div class="stat-box"><div class="stat-icon">🔥</div><div class="stat-val">{max(1, len(st.session_state.messages)//5)}</div><div class="stat-lbl">Daraja</div></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<p class="section-label">Funksiyalar</p>'
                '<p class="section-title">Nima qilmoqchisiz?</p>',
                unsafe_allow_html=True)

    st.markdown("""
    <div class="cards-grid">
        <div class="somo-card"><span class="card-icon">💬</span><div class="card-title">Chat AI</div><div class="card-desc">4 ta AI bilan aqlli suhbat</div></div>
        <div class="somo-card"><span class="card-icon">📊</span><div class="card-title">Excel Generator</div><div class="card-desc">Formulalar, ranglar, multi-sheet</div></div>
        <div class="somo-card"><span class="card-icon">📝</span><div class="card-title">Word Generator</div><div class="card-desc">Rezyume, shartnoma, biznes reja</div></div>
        <div class="somo-card"><span class="card-icon">💻</span><div class="card-title">Kod Generator</div><div class="card-desc">Python bot, API, ML model</div></div>
        <div class="somo-card"><span class="card-icon">🌐</span><div class="card-title">HTML Generator</div><div class="card-desc">Portfolio, landing, dashboard</div></div>
        <div class="somo-card"><span class="card-icon">📋</span><div class="card-title">CSV Generator</div><div class="card-desc">Dataset — bir so'rovda</div></div>
        <div class="somo-card"><span class="card-icon">🎨</span><div class="card-title">Shablonlar</div><div class="card-desc">16 ta tayyor shablon</div></div>
        <div class="somo-card"><span class="card-icon">🔍</span><div class="card-title">Hujjat Tahlili</div><div class="card-desc">PDF & DOCX tahlil</div></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)
    st.markdown('<p class="section-label">Tezkor Harakatlar</p>', unsafe_allow_html=True)

    q1, q2, q3, q4 = st.columns(4)
    for col, icon, label, page in [
        (q1, "📊", "Oylik Byudjet", "excel"),
        (q2, "📄", "Rezyume",       "word"),
        (q3, "🤖", "Telegram Bot",  "code"),
        (q4, "🌐", "Landing Page",  "html"),
    ]:
        with col:
            if st.button(f"{icon}  {label}", use_container_width=True, key=f"qk_{page}"):
                st.session_state.page = page
                st.rerun()

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: CHAT AI
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "chat":
    cur_prov = st.session_state.chat_provider

    st.markdown(f"""
    <div class="somo-hero">
        <div class="grid-dots"></div>
        <div class="somo-hero-content">
            <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#646cff;
                      margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
                ✦ Smart Chat · {api_badge(cur_prov)}
            </p>
            <h1>💬 Chat <span class="g-text">AI</span></h1>
            <p class="subtitle">So'zingizni yozing — AI tushunadi va javob beradi.</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Welcome cards (empty chat)
    if not st.session_state.messages:
        st.markdown("""
        <div class="cards-grid" style="grid-template-columns:repeat(auto-fill,minmax(175px,1fr));">
            <div class="somo-card"><span class="card-icon" style="font-size:26px;">📊</span><div class="card-title" style="font-size:12px;">"Oylik byudjet jadvali"</div><div class="card-desc">Excel avtomatik yaratiladi</div></div>
            <div class="somo-card"><span class="card-icon" style="font-size:26px;">📝</span><div class="card-title" style="font-size:12px;">"Rezyume yozing"</div><div class="card-desc">Word fayl tayyorlanadi</div></div>
            <div class="somo-card"><span class="card-icon" style="font-size:26px;">💻</span><div class="card-title" style="font-size:12px;">"Python kodi yozing"</div><div class="card-desc">.py fayl yuklab olish</div></div>
            <div class="somo-card"><span class="card-icon" style="font-size:26px;">🌐</span><div class="card-title" style="font-size:12px;">"Landing page yarat"</div><div class="card-desc">HTML fayl tayyorlanadi</div></div>
        </div>
        """, unsafe_allow_html=True)

    # Messages
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            if msg["role"] == "assistant" and "provider" in msg:
                st.markdown(
                    f'<div style="margin-bottom:8px;">{api_badge(msg.get("provider","groq"))}</div>',
                    unsafe_allow_html=True
                )
            st.markdown(msg["content"])
            if "file_data" in msg:
                fd = msg["file_data"]
                if fd.get("bytes"):
                    download_btn(fd["bytes"], fd["name"], fd["label"])

    # Document upload
    with st.expander("📂  Hujjat yuklash (PDF yoki DOCX)", expanded=False):
        upl = st.file_uploader("Fayl tanlang", type=["pdf", "docx"],
                               key="chat_upl", label_visibility="collapsed")
        if upl:
            with st.spinner("📄 O'qilmoqda..."):
                txt = process_document(upl)
                st.session_state.uploaded_text = txt
            if txt:
                st.success(f"✅ {upl.name} — {len(txt):,} belgi")
            else:
                st.error("❌ O'qilmadi")

    # Chat input
    if prompt := st.chat_input("💭  Yozing... Excel, Word, Kod, HTML so'rang!", key="chat_inp"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        db_log(uname, "User", prompt, "input", cur_prov)

        intent = detect_intent(prompt)

        with st.chat_message("assistant"):
            PAGE_MAP = {
                "excel": ("excel", "📊", "Excel Generator",  "Formulalar, ranglar, professional jadvallar"),
                "word":  ("word",  "📝", "Word Generator",   "Rezyume, shartnoma, biznes reja"),
                "code":  ("code",  "💻", "Kod Generator",    "Python bot, API, ML modeli"),
                "html":  ("html",  "🌐", "HTML Generator",   "Portfolio, landing page, dashboard"),
                "csv":   ("csv",   "📋", "CSV Generator",    "Dataset, ma'lumotlar to'plami"),
            }

            if intent in PAGE_MAP:
                _page, _em, _title, _desc = PAGE_MAP[intent]
                st.session_state[f"{_page}_prompt"] = prompt

                st.markdown(f"""
                <div style="background:linear-gradient(135deg,rgba(100,108,255,0.1),rgba(167,139,250,0.06));
                            border:1px solid rgba(100,108,255,0.25);border-radius:16px;padding:18px 20px;margin:8px 0;">
                    <div style="font-size:28px;margin-bottom:8px;">{_em}</div>
                    <div style="font-size:15px;font-weight:700;color:#f0f0ff;margin-bottom:4px;">{_title}</div>
                    <div style="font-size:12.5px;color:#a0a0c0;margin-bottom:14px;">{_desc}</div>
                    <div style="font-size:12px;color:#818cf8;background:rgba(100,108,255,0.08);
                                border-radius:8px;padding:8px 12px;font-family:'JetBrains Mono',monospace;">
                        📝 So'rovingiz saqlandi: "{prompt[:60]}{'...' if len(prompt)>60 else ''}"
                    </div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"{_em}  {_title} ga o'tish →",
                             key=f"goto_{_page}_{len(st.session_state.messages)}",
                             type="primary", use_container_width=True):
                    st.session_state.page = _page
                    st.rerun()

                st.session_state.messages.append({
                    "role": "assistant",
                    "content": f"{_em} **{_title}** da yaxshiroq natija olasiz!",
                    "provider": "none"
                })
            else:
                STYLES = {
                    "Aqlli yordamchi": "Sen Somo AI — aqlli, professional yordamchi. Foydalanuvchi tilida javob ber.",
                    "Do'stona":        "Sen Somo AI — do'stona, samimiy yordamchi. Foydalanuvchi tilida javob ber.",
                    "Rasmiy ekspert":  "Sen Somo AI — rasmiy, aniq ekspert. Foydalanuvchi tilida javob ber.",
                    "Ijodkor":         "Sen Somo AI — ijodkor, yangilikchi. Foydalanuvchi tilida javob ber.",
                    "Texnik":          "Sen Somo AI — chuqur texnik ekspert. Foydalanuvchi tilida javob ber.",
                }
                sys_txt = STYLES.get(st.session_state.ai_style, STYLES["Aqlli yordamchi"])

                msgs_ai = [{"role": "system", "content": sys_txt}]
                if st.session_state.uploaded_text:
                    msgs_ai.append({
                        "role": "system",
                        "content": f"Yuklangan hujjat:\n{st.session_state.uploaded_text[:4000]}"
                    })
                for m in st.session_state.messages[:-1][-20:]:
                    msgs_ai.append({"role": m["role"], "content": m["content"]})
                msgs_ai.append({"role": "user", "content": prompt})

                st.markdown(
                    f'<div style="margin-bottom:10px;">{api_badge(cur_prov)}</div>',
                    unsafe_allow_html=True
                )
                ph   = st.empty()
                resp = ""
                used = cur_prov

                for chunk, prov_name in call_ai_stream(
                    msgs_ai, st.session_state.temperature, provider=cur_prov
                ):
                    resp += chunk
                    used  = prov_name
                    ph.markdown(resp)

                ph.markdown(resp)
                db_log(uname, "Assistant", resp, "chat", used)
                st.session_state.messages.append({
                    "role": "assistant", "content": resp, "provider": used
                })

        st.rerun()

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: EXCEL
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "excel":
    st.markdown(f"""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#34d399;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ File Generator · {api_badge(st.session_state.excel_provider)}
        </p>
        <h1>📊 Excel <span class="g-text">Generator</span></h1>
        <p class="subtitle">Har qanday jadval, hisobot — AI bilan professional Excel faylga aylantiring.</p>
    </div></div>
    """, unsafe_allow_html=True)

    XL_EXAMPLES = [
        ("💰","Oylik Byudjet",    "12 oylik moliyaviy byudjet: daromad manbalari, xarajatlar, foyda, formulalar"),
        ("📦","Inventar",         "100 ta mahsulot: ID, nomi, kategoriya, miqdori, narxi, jami qiymat"),
        ("👥","Xodimlar Maoshi",  "Xodimlar ish haqi: ism, lavozim, maosh, bonus, soliq, sof maosh"),
        ("📈","Savdo Hisoboti",   "Oylik savdo: mahsulot, reja, haqiqat, farq, % bajarilish"),
        ("🎓","Talabalar Bahosi", "30 talaba 6 fandan baho: o'rtacha, reyting, davomat"),
        ("📅","Loyiha Jadvali",   "IT loyiha: vazifalar, mas'ul, boshlanish, tugash, holat, %"),
    ]
    c1, c2, c3 = st.columns(3)
    for i, (ico, title, fp) in enumerate(XL_EXAMPLES):
        with [c1, c2, c3][i % 3]:
            if st.button(f"{ico}  {title}", key=f"xlq_{i}", use_container_width=True):
                st.session_state["xl_prompt"] = fp
                st.rerun()

    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)
    col_i, col_o = st.columns([3, 1])

    with col_i:
        xl_prompt = st.text_area(
            "📝  Jadval tavsifi:",
            value=st.session_state.get("xl_prompt", st.session_state.get("excel_prompt", "")),
            placeholder="Masalan: 6 xodimlik IT kompaniya uchun ish haqi jadvali...",
            height=140, key="xl_in"
        )

    with col_o:
        provider_selector("excel_provider", "xl_prov")
        xl_temp = st.slider("Aniqlik", 0.0, 0.6, 0.15, 0.05, key="xl_temp")
        gen_xl  = st.button("🚀  Excel Yaratish", use_container_width=True,
                            type="primary", key="gen_xl")

    if gen_xl:
        if not xl_prompt.strip():
            st.warning("⚠️ Jadval tavsifini kiriting!")
        else:
            # Promptni tozalash (keyingi yuklaganda qolmasin)
            if "xl_prompt" in st.session_state: del st.session_state["xl_prompt"]
            if "excel_prompt" in st.session_state: del st.session_state["excel_prompt"]
            xl_prov = st.session_state.excel_provider
            st.markdown(
                f'<div class="somo-notify">📊 Excel yaratilmoqda... {api_badge(xl_prov)}</div>',
                unsafe_allow_html=True
            )
            prog = st.progress(0)
            for pct in range(0, 75, 12):
                time.sleep(0.2); prog.progress(pct)
            fb, fn = gen_excel(xl_prompt, xl_temp, provider=xl_prov)
            prog.progress(100); time.sleep(0.1); prog.empty()
            if fb and isinstance(fb, bytes):
                st.session_state.files_count += 1
                download_btn(fb, fn, "Excel")
            else:
                st.error(f"❌ {fn}")

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: WORD
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "word":
    st.markdown(f"""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#38bdf8;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ File Generator · {api_badge(st.session_state.word_provider)}
        </p>
        <h1>📝 Word <span class="g-text">Generator</span></h1>
        <p class="subtitle">Professional hujjatlar — rezyume, shartnoma, biznes reja — AI bilan bir daqiqada.</p>
    </div></div>
    """, unsafe_allow_html=True)

    WD_EXAMPLES = [
        ("👤","Rezyume / CV",       "Python backend dasturchi rezyume: 4 yil tajriba, ko'nikmalar, ta'lim, sertifikatlar"),
        ("🤝","Hamkorlik Xati",     "IT kompaniyalar hamkorlik taklifnomasi: taqdimot, taklif, foyda, shartlar"),
        ("📋","Ijara Shartnomasi",  "Turar joy ijara shartnomasi: tomonlar, ob'ekt, muddat, to'lov, mas'uliyat"),
        ("📖","Biznes Reja",        "Startap biznes reja: bozor tahlili, mahsulot, marketing, moliyaviy prognoz"),
        ("🎓","Kurs Ishi",          "Sun'iy intellekt kurs ishi: 3 bob, xulosa, adabiyotlar, 15+ sahifa"),
        ("📑","Buyruq",             "Kompaniya direktori buyrug'i: xodim qabul, lavozim, ish haqi"),
    ]
    c1, c2, c3 = st.columns(3)
    for i, (ico, title, fp) in enumerate(WD_EXAMPLES):
        with [c1, c2, c3][i % 3]:
            if st.button(f"{ico}  {title}", key=f"wdq_{i}", use_container_width=True):
                st.session_state["wd_prompt"] = fp
                st.rerun()

    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)
    col_i, col_o = st.columns([3, 1])

    with col_i:
        wd_prompt = st.text_area(
            "📝  Hujjat tavsifi:",
            value=st.session_state.get("wd_prompt", st.session_state.get("word_prompt", "")),
            placeholder="Masalan: IT kompaniya uchun dasturchi mehnat shartnomasi...",
            height=140, key="wd_in"
        )

    with col_o:
        provider_selector("word_provider", "wd_prov")
        gen_wd = st.button("🚀  Word Yaratish", use_container_width=True,
                           type="primary", key="gen_wd")

    if gen_wd:
        if not wd_prompt.strip():
            st.warning("⚠️ Hujjat tavsifini kiriting!")
        else:
            if "wd_prompt" in st.session_state: del st.session_state["wd_prompt"]
            if "word_prompt" in st.session_state: del st.session_state["word_prompt"]
            wd_prov = st.session_state.word_provider
            st.markdown(
                f'<div class="somo-notify">📝 Word yaratilmoqda... {api_badge(wd_prov)}</div>',
                unsafe_allow_html=True
            )
            prog = st.progress(0)
            for pct in range(0, 75, 15):
                time.sleep(0.2); prog.progress(pct)
            fb, fn = gen_word(wd_prompt, provider=wd_prov)
            prog.progress(100); time.sleep(0.1); prog.empty()
            if fb and isinstance(fb, bytes):
                st.session_state.files_count += 1
                download_btn(fb, fn, "Word")
            else:
                st.error(f"❌ {fn}")

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: CODE
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "code":
    st.markdown(f"""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#fbbf24;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ Code Generator · {api_badge(st.session_state.code_provider)}
        </p>
        <h1>💻 Kod <span class="g-text">Generator</span></h1>
        <p class="subtitle">Professional Python kodi — error handling, izohlar va best practices bilan.</p>
    </div></div>
    """, unsafe_allow_html=True)

    CD_EXAMPLES = [
        ("🤖","Telegram Bot",    "Aiogram v3 Telegram bot: /start, /help, inline keyboard, FSM, SQLite, .env"),
        ("🌐","FastAPI CRUD",    "FastAPI CRUD API: PostgreSQL, SQLAlchemy, Pydantic, JWT auth, Swagger"),
        ("📊","Streamlit App",   "Streamlit dashboard: CSV yuklash, pandas, plotly grafiklar, filterlar"),
        ("🔍","Web Scraper",     "BeautifulSoup4 scraper: sahifa tahlili, CSV saqlash, delay, headers"),
        ("🤖","ML Model",        "Scikit-learn classification: preprocessing, train, hyperparameter, report"),
        ("📧","Email Sender",    "smtplib email: HTML template, attachment, bulk send, logging"),
    ]
    c1, c2, c3 = st.columns(3)
    for i, (ico, title, fp) in enumerate(CD_EXAMPLES):
        with [c1, c2, c3][i % 3]:
            if st.button(f"{ico}  {title}", key=f"cdq_{i}", use_container_width=True):
                st.session_state["cd_prompt"] = fp
                st.rerun()

    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)
    col_i, col_o = st.columns([3, 1])

    with col_i:
        cd_prompt = st.text_area(
            "📝  Kod tavsifi:",
            value=st.session_state.get("cd_prompt", st.session_state.get("code_prompt", "")),
            placeholder="Masalan: Telegram bot — narx so'raganda Olx.uz dan avtomatik qidirsin...",
            height=140, key="cd_in"
        )

    with col_o:
        provider_selector("code_provider", "cd_prov")
        cd_temp = st.slider("Ijodkorlik", 0.0, 0.5, 0.1, 0.05, key="cd_temp")
        gen_cd  = st.button("🚀  Kod Yaratish", use_container_width=True,
                            type="primary", key="gen_cd")

    if gen_cd:
        if not cd_prompt.strip():
            st.warning("⚠️ Kod tavsifini kiriting!")
        else:
            if "cd_prompt" in st.session_state: del st.session_state["cd_prompt"]
            if "code_prompt" in st.session_state: del st.session_state["code_prompt"]
            cd_prov = st.session_state.code_provider
            st.markdown(
                f'<div class="somo-notify">💻 Kod yozilmoqda... {api_badge(cd_prov)}</div>',
                unsafe_allow_html=True
            )
            prog = st.progress(0)
            for pct in range(0, 65, 15):
                time.sleep(0.18); prog.progress(pct)
            fb, fn = gen_code(cd_prompt, cd_temp, provider=cd_prov)
            prog.progress(100); prog.empty()
            st.session_state.files_count += 1
            st.markdown('<div class="somo-success">✅ Kod tayyor!</div>', unsafe_allow_html=True)
            with st.expander("👁  Kod Preview", expanded=True):
                st.code(fb.decode("utf-8"), language="python")
            st.download_button(
                f"⬇️  {fn}", fb, fn, "text/x-python",
                use_container_width=True, type="primary",
                key=f"dl_py_{time.time()}"
            )

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: HTML
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "html":
    st.markdown(f"""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#f472b6;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ HTML Generator · {api_badge(st.session_state.html_provider)}
        </p>
        <h1>🌐 HTML <span class="g-text">Generator</span></h1>
        <p class="subtitle">Zamonaviy, animatsiyali veb sahifalar — bitta .html faylda hamma narsa.</p>
    </div></div>
    """, unsafe_allow_html=True)

    HT_EXAMPLES = [
        ("🎨","Portfolio",        "Web developer portfolio: hero typewriter, skills bars, projects, dark glassmorphism"),
        ("🛒","Mahsulot Sahifa",  "E-commerce: gallery, narx, add to cart, reviews, responsive"),
        ("📊","Dashboard",        "Analytics dashboard: sidebar, stat cards, Chart.js grafiklari"),
        ("🎪","Event Landing",    "Konferensiya: hero countdown, speakers, schedule, tickets CTA"),
        ("🔐","Login Sahifa",     "Auth UI: glassmorphism login/register, validation, particles.js"),
        ("📰","Blog Post",        "Blog: hero image, typography, sidebar, TOC, code highlight"),
    ]
    c1, c2, c3 = st.columns(3)
    for i, (ico, title, fp) in enumerate(HT_EXAMPLES):
        with [c1, c2, c3][i % 3]:
            if st.button(f"{ico}  {title}", key=f"htq_{i}", use_container_width=True):
                st.session_state["ht_prompt"] = fp
                st.rerun()

    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)
    col_i, col_o = st.columns([3, 1])

    with col_i:
        ht_prompt = st.text_area(
            "📝  Sahifa tavsifi:",
            value=st.session_state.get("ht_prompt", st.session_state.get("html_prompt", "")),
            placeholder="Masalan: AI kompaniyasi uchun landing page — dark neon dizayn...",
            height=140, key="ht_in"
        )

    with col_o:
        provider_selector("html_provider", "ht_prov")
        gen_ht = st.button("🚀  HTML Yaratish", use_container_width=True,
                           type="primary", key="gen_ht")

    if gen_ht:
        if not ht_prompt.strip():
            st.warning("⚠️ Sahifa tavsifini kiriting!")
        else:
            if "ht_prompt" in st.session_state: del st.session_state["ht_prompt"]
            if "html_prompt" in st.session_state: del st.session_state["html_prompt"]
            ht_prov = st.session_state.html_provider
            st.markdown(
                f'<div class="somo-notify">🌐 HTML yaratilmoqda... {api_badge(ht_prov)}</div>',
                unsafe_allow_html=True
            )
            prog = st.progress(0)
            for pct in range(0, 70, 14):
                time.sleep(0.2); prog.progress(pct)
            fb, fn = gen_html(ht_prompt, 0.5, provider=ht_prov)
            prog.progress(100); prog.empty()
            st.session_state.files_count += 1
            st.markdown('<div class="somo-success">✅ HTML tayyor! Yuklab, brauzerda oching.</div>',
                        unsafe_allow_html=True)
            with st.expander("👁  HTML Preview"):
                st.code(fb.decode("utf-8")[:3000], language="html")
            st.download_button(
                f"⬇️  {fn}", fb, fn, "text/html",
                use_container_width=True, type="primary",
                key=f"dl_html_{time.time()}"
            )

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: CSV
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "csv":
    st.markdown(f"""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#a78bfa;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ Data Generator · {api_badge(st.session_state.csv_provider)}
        </p>
        <h1>📋 CSV <span class="g-text">Generator</span></h1>
        <p class="subtitle">Katta ma'lumotlar to'plamini bir so'rovda yarating.</p>
    </div></div>
    """, unsafe_allow_html=True)

    CV_EXAMPLES = [
        ("📦","Mahsulotlar",       "100 ta mahsulot: ID, nomi, kategoriya, narxi, miqdori, brend, reyting"),
        ("👥","Foydalanuvchilar",  "50 ta user: ID, ism, email, telefon, shahar, ro'yxat sanasi, holat"),
        ("🌍","Mamlakatlar",       "Dunyo mamlakatlari: nomi, poytaxti, aholisi, maydoni, YIM, hudud"),
        ("📱","Ilovalar",          "Top 100 mobil ilova: nomi, kategoriya, reyting, yuklamalar, narxi"),
        ("🎬","Filmlar",           "Top 100 film: nomi, rejissor, yili, janri, IMDB reyting, byudjet"),
        ("💼","Kompaniyalar",      "50 kompaniya: nomi, sektori, xodimlar soni, daromad, asos yili"),
    ]
    c1, c2, c3 = st.columns(3)
    for i, (ico, title, fp) in enumerate(CV_EXAMPLES):
        with [c1, c2, c3][i % 3]:
            if st.button(f"{ico}  {title}", key=f"cvq_{i}", use_container_width=True):
                st.session_state["cv_prompt"] = fp
                st.rerun()

    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)
    col_i, col_o = st.columns([3, 1])

    with col_i:
        cv_prompt = st.text_area(
            "📝  Dataset tavsifi:",
            value=st.session_state.get("cv_prompt", ""),
            placeholder="Masalan: 80 ta O'zbekiston shahri: viloyati, aholisi, maydoni...",
            height=130, key="cv_in"
        )

    with col_o:
        provider_selector("csv_provider", "cv_prov")
        gen_cv = st.button("🚀  CSV Yaratish", use_container_width=True,
                           type="primary", key="gen_cv")

    if gen_cv:
        if not cv_prompt.strip():
            st.warning("⚠️ Dataset tavsifini kiriting!")
        else:
            if "cv_prompt" in st.session_state: del st.session_state["cv_prompt"]
            cv_prov = st.session_state.csv_provider
            st.markdown(
                f'<div class="somo-notify">📋 Dataset yaratilmoqda... {api_badge(cv_prov)}</div>',
                unsafe_allow_html=True
            )
            prog = st.progress(0)
            for pct in range(0, 65, 15):
                time.sleep(0.18); prog.progress(pct)
            fb, fn = gen_csv(cv_prompt, provider=cv_prov)
            prog.progress(100); prog.empty()
            st.session_state.files_count += 1
            try:
                df = pd.read_csv(io.BytesIO(fb))
                st.markdown(
                    f'<div class="somo-success">✅ CSV tayyor — {len(df)} satr, {len(df.columns)} ustun</div>',
                    unsafe_allow_html=True
                )
                st.dataframe(df.head(10), use_container_width=True)
                if len(df) > 10:
                    st.caption(f"↑ Birinchi 10 ta (jami {len(df)} ta)")
            except Exception:
                st.markdown('<div class="somo-success">✅ CSV tayyor!</div>', unsafe_allow_html=True)
            st.download_button(
                f"⬇️  {fn}", fb, fn, "text/csv",
                use_container_width=True, type="primary",
                key=f"dl_csv_{time.time()}"
            )

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: TEMPLATES
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "templates":
    st.markdown("""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#646cff;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ Template Library
        </p>
        <h1>🎨 Shablonlar <span class="g-text">Markazi</span></h1>
        <p class="subtitle">16 ta professional shablon — bitta bosish bilan yarating.</p>
    </div></div>
    """, unsafe_allow_html=True)

    TEMPLATES: Dict[str, List[Dict]] = {
        "📊 Biznes": [
            {"ico":"💰","title":"Oylik Byudjet",   "tag":"excel","tag_cls":"tag-excel","desc":"12 oylik moliyaviy byudjet",
             "prompt":"12 oylik moliyaviy byudjet Excel: daromad, xarajatlar 8 kategoriya, sof foyda, jamg'arma, formulalar"},
            {"ico":"📈","title":"KPI Dashboard",   "tag":"excel","tag_cls":"tag-excel","desc":"Kompaniya KPI ko'rsatkichlari",
             "prompt":"Kompaniya KPI Excel: 15 ko'rsatkich, maqsad va haqiqat, farq foizi, RAG ranglash"},
            {"ico":"📋","title":"Biznes Reja",     "tag":"word", "tag_cls":"tag-word", "desc":"To'liq startap biznes reja",
             "prompt":"IT startap biznes reja Word: ijroiya xulosa, bozor tahlili, mahsulot, marketing, moliyaviy prognoz 3 yil"},
            {"ico":"🤝","title":"Hamkorlik Xati",  "tag":"word", "tag_cls":"tag-word", "desc":"Professional hamkorlik taklifnomasi",
             "prompt":"Hamkorlik taklifnomasi Word: kompaniya taqdimoti, taklif, o'zaro foyda, shartlar"},
        ],
        "💻 Dasturlash": [
            {"ico":"🤖","title":"Telegram Bot",    "tag":"code", "tag_cls":"tag-code", "desc":"Aiogram v3, FSM, SQLite",
             "prompt":"Aiogram v3 Telegram bot: /start, /help, InlineKeyboard, FSM, SQLite, admin panel, .env"},
            {"ico":"🌐","title":"FastAPI REST",    "tag":"code", "tag_cls":"tag-code", "desc":"CRUD, JWT, PostgreSQL",
             "prompt":"FastAPI REST API: User, Post, Comment, SQLAlchemy+PostgreSQL, Pydantic, JWT, CRUD, CORS"},
            {"ico":"🎨","title":"Portfolio Sayt",  "tag":"html", "tag_cls":"tag-html", "desc":"Dark theme, glassmorphism",
             "prompt":"Web developer portfolio HTML: typewriter hero, skills bars, projects glassmorphism, dark theme, responsive"},
            {"ico":"📊","title":"Streamlit App",   "tag":"code", "tag_cls":"tag-code", "desc":"Data dashboard, grafik",
             "prompt":"Streamlit data app: CSV yuklab, pandas, Plotly grafiklar, filterlar, dark theme"},
        ],
        "📚 Ta'lim": [
            {"ico":"📖","title":"Dars Rejasi",     "tag":"word", "tag_cls":"tag-word", "desc":"45 daqiqalik dars konspekti",
             "prompt":"Python asoslari 45 daqiqalik dars Word: maqsadlar, bosqichlar, savol-javob, baholash"},
            {"ico":"📝","title":"Test Savollari",  "tag":"excel","tag_cls":"tag-excel","desc":"25 ta test, 4 variant",
             "prompt":"Python 25 test Excel: savol, A-B-C-D variant, to'g'ri javob, mavzu, qiyinchilik"},
            {"ico":"🎓","title":"Baholash Jadvali","tag":"excel","tag_cls":"tag-excel","desc":"30 talaba, 6 fan",
             "prompt":"Guruh baholash Excel: 30 talaba, 6 fan, o'rtacha, GPA, reyting, davomat, formulalar"},
            {"ico":"📚","title":"Kurs Ishi",       "tag":"word", "tag_cls":"tag-word", "desc":"15+ sahifa, 3 bob",
             "prompt":"Sun'iy intellekt kurs ishi Word: titul, mundarija, kirish, 3 bob, xulosa, 15 adabiyot"},
        ],
        "👤 Shaxsiy": [
            {"ico":"📄","title":"Rezyume / CV",    "tag":"word", "tag_cls":"tag-word", "desc":"Professional CV",
             "prompt":"Python/Django dasturchi rezyume Word: ism, kontakt, xulosa, ko'nikmalar, 2 ish joyi, ta'lim, sertifikatlar"},
            {"ico":"📅","title":"Haftalik Reja",   "tag":"excel","tag_cls":"tag-excel","desc":"7 kun, vazifalar, holat",
             "prompt":"Haftalik vazifalar Excel: 7 kun, vaqt slotlari, vazifa, ustuvorlik, holat, statistika"},
            {"ico":"💰","title":"Shaxsiy Byudjet", "tag":"excel","tag_cls":"tag-excel","desc":"Daromad, xarajat, jamg'arma",
             "prompt":"Shaxsiy moliya Excel: daromad, xarajatlar kategoriyalar, jamg'arma, oylik trend"},
            {"ico":"💪","title":"Sport Rejasi",    "tag":"excel","tag_cls":"tag-excel","desc":"3 oylik trening, progres",
             "prompt":"3 oylik sport Excel: haftalik trening, mashqlar, to'plamlar, ozish maqsad, kaloriya, progres"},
        ],
    }

    sel_cat = st.selectbox("📁  Kategoriya:", list(TEMPLATES.keys()), key="tmpl_sel")
    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)

    items = TEMPLATES[sel_cat]
    tc1, tc2 = st.columns(2)

    TAG_TO_PROV = {
        "excel": st.session_state.excel_provider,
        "word":  st.session_state.word_provider,
        "code":  st.session_state.code_provider,
        "html":  st.session_state.html_provider,
    }

    for i, tmpl in enumerate(items):
        with [tc1, tc2][i % 2]:
            t_prov = TAG_TO_PROV.get(tmpl["tag"], "groq")
            st.markdown(f"""
            <div class="tmpl-card">
                <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:8px;">
                    <span class="tmpl-tag {tmpl['tag_cls']}">{tmpl['tag'].upper()}</span>
                    {api_badge(t_prov)}
                </div>
                <div class="tmpl-title">{tmpl['ico']}  {tmpl['title']}</div>
                <div class="tmpl-desc">{tmpl['desc']}</div>
            </div>
            """, unsafe_allow_html=True)

            bc1, bc2 = st.columns(2)
            with bc1:
                if st.button("🚀  Yaratish", key=f"tgen_{sel_cat}_{i}",
                             use_container_width=True, type="primary"):
                    with st.spinner("⏳ Tayyorlanmoqda..."):
                        tag = tmpl["tag"]
                        if tag == "excel":
                            fb, fn = gen_excel(tmpl["prompt"], provider=st.session_state.excel_provider)
                            if fb and isinstance(fb, bytes):
                                st.session_state.files_count += 1
                                st.download_button("⬇️  Excel", fb, fn, MIME_TYPES["xlsx"],
                                                   key=f"tdl_{sel_cat}_{i}", type="primary")
                        elif tag == "word":
                            fb, fn = gen_word(tmpl["prompt"], provider=st.session_state.word_provider)
                            if fb and isinstance(fb, bytes):
                                st.session_state.files_count += 1
                                st.download_button("⬇️  Word", fb, fn, MIME_TYPES["docx"],
                                                   key=f"tdl_{sel_cat}_{i}", type="primary")
                        elif tag == "code":
                            fb, fn = gen_code(tmpl["prompt"], provider=st.session_state.code_provider)
                            st.session_state.files_count += 1
                            st.download_button("⬇️  .py", fb, fn, MIME_TYPES["py"],
                                               key=f"tdl_{sel_cat}_{i}", type="primary")
                        elif tag == "html":
                            fb, fn = gen_html(tmpl["prompt"], provider=st.session_state.html_provider)
                            st.session_state.files_count += 1
                            st.download_button("⬇️  HTML", fb, fn, MIME_TYPES["html"],
                                               key=f"tdl_{sel_cat}_{i}", type="primary")
            with bc2:
                if st.button("💬  Chat AI", key=f"tchat_{sel_cat}_{i}",
                             use_container_width=True):
                    st.session_state.messages.append({
                        "role": "user", "content": tmpl["prompt"]
                    })
                    st.session_state.page = "chat"
                    st.rerun()

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: ANALYZE
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "analyze":
    st.markdown(f"""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#22d3ee;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ Document Analysis · {api_badge(st.session_state.analyze_provider)}
        </p>
        <h1>🔍 Hujjat <span class="g-text">Tahlili</span></h1>
        <p class="subtitle">PDF yoki Word faylni yuklang — AI xulosa, g'oyalar, savol-javob.</p>
    </div></div>
    """, unsafe_allow_html=True)

    col_up, col_act = st.columns([1, 1])

    with col_up:
        st.markdown('<p class="section-label">Fayl Yuklash</p>', unsafe_allow_html=True)
        provider_selector("analyze_provider", "az_prov", "🤖 AI Provider")

        upl_az = st.file_uploader("PDF yoki DOCX", type=["pdf", "docx"],
                                  key="az_upl", label_visibility="collapsed")
        if upl_az:
            with st.spinner("📄 O'qilmoqda..."):
                txt = process_document(upl_az)
                st.session_state.uploaded_text = txt
            if txt:
                st.markdown(
                    f'<div class="somo-success">✅ {upl_az.name} — {len(txt):,} belgi</div>',
                    unsafe_allow_html=True
                )
                with st.expander("👁  Matnni ko'rish"):
                    st.text(txt[:2000] + ("..." if len(txt) > 2000 else ""))
            else:
                st.error("❌ Fayl o'qilmadi")

    with col_act:
        st.markdown('<p class="section-label">Tahlil Amaliyotlari</p>', unsafe_allow_html=True)

        if st.session_state.uploaded_text:
            az_prov = st.session_state.analyze_provider
            ACTIONS = {
                "📝  Qisqa Xulosa":    "Hujjatni 5-7 asosiy band bilan xulosala. Har bandni ★ bilan boshlat.",
                "🔑  Kalit G'oyalar":  "Hujjatdagi 8-10 muhim g'oya va faktlarni ro'yxat shaklida ajrat.",
                "❓  Savol-Javob":     "Hujjat bo'yicha 10 muhim savol tuz va har biriga javob ber.",
                "🌐  Inglizcha":       "Hujjat mazmunini professional ingliz tiliga tarjima qil.",
                "📊  Statistika":      "Hujjatdagi barcha raqamlar va statistikani jadval ko'rinishida tizimlashtir.",
                "✅  Action Items":    "Hujjatdan aniq amaliy vazifalar va keyingi qadamlarni ustuvorlik bo'yicha tartibla.",
            }
            for act_lbl, act_prompt in ACTIONS.items():
                if st.button(act_lbl, key=f"az_{act_lbl}", use_container_width=True):
                    az_msgs = [
                        {"role": "system", "content": "Sen professional hujjat tahlilchisan. Foydalanuvchi tilida javob ber."},
                        {"role": "user",   "content": f"Hujjat:\n{st.session_state.uploaded_text[:4500]}\n\nVazifa: {act_prompt}"},
                    ]
                    st.markdown(f'<div style="margin-bottom:6px;">{api_badge(az_prov)}</div>',
                                unsafe_allow_html=True)
                    az_ph = st.empty()
                    az_full = ""
                    for chunk, _ in call_ai_stream(az_msgs, temperature=0.4, provider=az_prov):
                        az_full += chunk
                        az_ph.markdown(f"**{act_lbl}**\n\n{az_full}")
                    az_ph.markdown(f"**{act_lbl}**\n\n{az_full}")
        else:
            st.markdown("""
            <div style="background:rgba(100,108,255,0.04);border:1px solid rgba(100,108,255,0.12);
                       border-radius:14px;padding:28px;text-align:center;margin-top:8px;">
                <p style="font-size:36px;margin-bottom:12px;">📂</p>
                <p style="color:#50506a;font-size:14px;">Chap tomonda fayl yuklang</p>
            </div>
            """, unsafe_allow_html=True)

    if st.session_state.uploaded_text:
        st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)
        st.markdown('<p class="section-label">O\'z Savolingiz</p>', unsafe_allow_html=True)
        cq_col, cb_col = st.columns([4, 1])
        with cq_col:
            custom_q = st.text_input("", placeholder="🔍 Hujjat haqida savolingiz...",
                                     label_visibility="collapsed", key="az_q")
        with cb_col:
            if st.button("🔍  Qidirish", use_container_width=True,
                         type="primary", key="az_ask"):
                if custom_q.strip():
                    az_prov = st.session_state.analyze_provider
                    az_msgs = [
                        {"role": "system", "content": "Hujjat asosida aniq javob ber."},
                        {"role": "user",   "content": f"Hujjat:\n{st.session_state.uploaded_text[:4500]}\n\nSavol: {custom_q}"},
                    ]
                    ans_ph  = st.empty()
                    ans_txt = ""
                    for chunk, _ in call_ai_stream(az_msgs, temperature=0.3, provider=az_prov):
                        ans_txt += chunk
                        ans_ph.markdown(f"**💬 Javob:**\n\n{ans_txt}")
                    ans_ph.markdown(f"**💬 Javob:**\n\n{ans_txt}")

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: HISTORY
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "history":
    st.markdown("""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#646cff;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ History
        </p>
        <h1>📜 Chat <span class="g-text">Tarixi</span></h1>
        <p class="subtitle">Barcha suhbatlaringiz — qidirish va eksport bilan.</p>
    </div></div>
    """, unsafe_allow_html=True)

    msgs = st.session_state.messages

    if msgs:
        u_cnt = sum(1 for m in msgs if m["role"] == "user")
        a_cnt = len(msgs) - u_cnt

        st.markdown(f"""
        <div class="stat-row" style="grid-template-columns:repeat(4,1fr);">
            <div class="stat-box"><div class="stat-icon">💬</div><div class="stat-val">{len(msgs)}</div><div class="stat-lbl">Jami</div></div>
            <div class="stat-box"><div class="stat-icon">👤</div><div class="stat-val">{u_cnt}</div><div class="stat-lbl">Sizdan</div></div>
            <div class="stat-box"><div class="stat-icon">🤖</div><div class="stat-val">{a_cnt}</div><div class="stat-lbl">AI dan</div></div>
            <div class="stat-box"><div class="stat-icon">📁</div><div class="stat-val">{st.session_state.files_count}</div><div class="stat-lbl">Fayllar</div></div>
        </div>
        """, unsafe_allow_html=True)

        col_s, col_e1, col_e2 = st.columns([3, 1, 1])
        with col_s:
            search = st.text_input("", placeholder="🔍 Xabarlarda qidirish...",
                                   label_visibility="collapsed", key="hist_s")

        safe_msgs = [{"role": m.get("role",""), "content": m.get("content",""),
                      "provider": m.get("provider","")} for m in msgs]
        txt_exp = "\n\n".join([f"[{m['role'].upper()}]\n{m['content']}" for m in msgs])

        with col_e1:
            st.download_button(
                "📥  JSON",
                json.dumps(safe_msgs, ensure_ascii=False, indent=2).encode(),
                f"somo_{datetime.now():%Y%m%d}.json",
                use_container_width=True
            )
        with col_e2:
            st.download_button(
                "📄  TXT", txt_exp.encode(),
                f"somo_{datetime.now():%Y%m%d}.txt",
                use_container_width=True
            )

        show = [m for m in msgs
                if search.lower() in m.get("content","").lower()] if search else msgs

        if search:
            st.markdown(
                f'<div class="somo-notify">🔍 "{search}" — {len(show)} ta natija</div>',
                unsafe_allow_html=True
            )

        st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)

        for msg in reversed(show[-50:]):
            is_user  = msg["role"] == "user"
            prov_u   = msg.get("provider", "groq")
            role_lbl = "👤  Siz" if is_user else f"🤖  Somo AI · {API_CONFIGS.get(prov_u, API_CONFIGS['groq'])['icon']}"
            body     = msg.get("content","")[:350] + ("..." if len(msg.get("content","")) > 350 else "")
            st.markdown(f"""
            <div class="hist-msg {'hist-user' if is_user else 'hist-ai'}">
                <div class="hist-role">{role_lbl}</div>
                <div class="hist-body">{body}</div>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="text-align:center;padding:60px 20px;">
            <div style="font-size:56px;margin-bottom:20px;">💬</div>
            <p style="color:#50506a;font-size:18px;font-weight:700;">Chat tarixi yo'q</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("💬  Chat AI ga o'tish", type="primary", key="hist_goto"):
            st.session_state.page = "chat"
            st.rerun()

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: FEEDBACK
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "feedback":
    st.markdown("""
    <div class="somo-hero"><div class="grid-dots"></div><div class="somo-hero-content">
        <p style="font-size:11px;letter-spacing:3.5px;font-weight:700;color:#f472b6;
                  margin-bottom:10px;text-transform:uppercase;font-family:'JetBrains Mono',monospace;">
            ✦ Feedback
        </p>
        <h1>💌 Fikr <span class="g-text">Bildirish</span></h1>
        <p class="subtitle">Sizning fikringiz Somo AI ni yaxshiroq qilishga yordam beradi.</p>
    </div></div>
    """, unsafe_allow_html=True)

    col_f, col_s = st.columns([3, 2])

    with col_f:
        with st.form("fb_form", clear_on_submit=True):
            rating = st.select_slider(
                "⭐  Baho:",
                options=[1, 2, 3, 4, 5], value=5,
                format_func=lambda x: "⭐" * x + f"  ({x}/5)"
            )
            category = st.selectbox("📂  Kategoriya:", [
                "Umumiy fikr",
                "Xato haqida",
                "Yangi funksiya taklifi",
                "Dizayn taklifi",
                "Tezlik muammosi",
                "Boshqa",
            ])
            message = st.text_area("✍️  Xabar:", height=140,
                                   placeholder="Fikrlaringizni yozing (kamida 10 belgi)...")
            email = st.text_input("📧  Email (ixtiyoriy):",
                                  placeholder="javob olish uchun")
            sub_fb = st.form_submit_button("📤  Yuborish",
                                           use_container_width=True, type="primary")
            if sub_fb:
                if not message or len(message) < 10:
                    st.error("❌ Kamida 10 belgidan iborat xabar yozing!")
                elif fb_db:
                    try:
                        fb_db.append_row([
                            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            st.session_state.username,
                            rating, category, message,
                            email or "N/A", "Yangi",
                            st.session_state.files_count
                        ])
                        st.balloons()
                        st.success("✅ Rahmat! Fikringiz yuborildi 🙏")
                    except Exception as e:
                        st.error(f"❌ {e}")
                else:
                    st.error("❌ Baza ulanmagan")

    with col_s:
        st.markdown("""
        <div style="background:var(--bg-card);border:1px solid var(--border);
                   border-radius:16px;padding:24px;margin-top:8px;">
            <p style="font-size:28px;text-align:center;margin-bottom:16px;">💬</p>
            <p style="font-size:14px;font-weight:700;color:#f0f0ff;margin-bottom:8px;">Nima haqida yozsam?</p>
            <ul style="color:#a0a0c0;font-size:13px;line-height:2;padding-left:16px;">
                <li>Fayllar to'g'ri chiqmadi</li>
                <li>Qo'shimcha funksiya istaklari</li>
                <li>Dizayn takliflari</li>
                <li>AI javoblari sifati</li>
                <li>Tezlik va ishlash muammosi</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════════════════
# PAGE: PROFILE
# ════════════════════════════════════════════════════════════════════════════════

elif st.session_state.page == "profile":
    mins  = (datetime.now() - st.session_state.login_time).seconds // 60
    avail = [p for p in PROVIDER_ORDER if p in ai_clients]

    st.markdown(f"""
    <div class="somo-hero" style="text-align:center;">
        <div class="grid-dots"></div>
        <div class="somo-hero-content">
            <div style="width:80px;height:80px;
                       background:linear-gradient(135deg,#4f46e5,#7c3aed,#f472b6);
                       border-radius:24px;margin:0 auto 18px;
                       display:flex;align-items:center;justify-content:center;
                       font-size:36px;font-weight:900;color:white;
                       box-shadow:0 0 40px rgba(100,108,255,0.5);">
                {uname[0].upper()}
            </div>
            <h1 style="font-size:28px;">{uname}</h1>
            <p style="color:rgba(255,255,255,0.5);font-size:13px;margin-top:8px;">
                🟢 ONLINE · Somo AI v4.0 · {len(avail)} API aktiv
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    p_stats = [
        ("💬", len(st.session_state.messages), "Xabarlar"),
        ("📁", st.session_state.files_count,    "Fayllar"),
        ("⏱",  mins,                            "Daqiqa"),
        ("🤖", len(avail),                      "API"),
    ]
    cols_ps = st.columns(4)
    for col, (icon, val, lbl) in zip(cols_ps, p_stats):
        with col:
            st.markdown(f"""
            <div class="profile-stat">
                <div style="font-size:24px;margin-bottom:8px;">{icon}</div>
                <div class="stat-val">{val}</div>
                <div class="stat-lbl">{lbl}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown('<hr class="somo-divider">', unsafe_allow_html=True)

    col_pw, col_info = st.columns(2)

    with col_pw:
        st.markdown('#### 🔑 Parolni O\'zgartirish')
        with st.form("pw_form"):
            old_pw  = st.text_input("Joriy parol",  type="password", key="pw_old")
            new_pw  = st.text_input("Yangi parol",  type="password", key="pw_new",
                                    placeholder="Kamida 6 belgi")
            conf_pw = st.text_input("Tasdiqlash",   type="password", key="pw_conf")
            if st.form_submit_button("🔄  Yangilash",
                                     type="primary", use_container_width=True):
                if len(new_pw) < 6:
                    st.error("❌ Yangi parol kamida 6 belgi!")
                elif new_pw != conf_pw:
                    st.error("❌ Parollar mos emas!")
                elif user_db:
                    try:
                        recs = get_all_users()
                        match = next(
                            (i for i, r in enumerate(recs)
                             if str(r.get("username")) == uname
                             and verify_password(old_pw, str(r.get("password", "")))),
                            None
                        )
                        if match is not None:
                            user_db.update_cell(match + 2, 2, hash_password(new_pw))
                            get_all_users.clear()
                            st.success("✅ Parol yangilandi!")
                        else:
                            st.error("❌ Joriy parol noto'g'ri!")
                    except Exception as e:
                        st.error(f"❌ {e}")
                else:
                    st.error("❌ Baza ulanmagan")

    with col_info:
        st.markdown('#### 🤖 Har Format uchun AI')
        prov_labels_all = [f"{API_CONFIGS[p]['icon']} {API_CONFIGS[p]['name']}" for p in avail]

        FORMAT_SETTINGS = [
            ("chat_provider",    "💬 Chat"),
            ("excel_provider",   "📊 Excel"),
            ("word_provider",    "📝 Word"),
            ("code_provider",    "💻 Kod"),
            ("html_provider",    "🌐 HTML"),
            ("csv_provider",     "📋 CSV"),
            ("analyze_provider", "🔍 Tahlil"),
        ]

        for sess_key, label in FORMAT_SETTINGS:
            curr = st.session_state.get(sess_key, "groq")
            cidx = avail.index(curr) if curr in avail else 0
            sel  = st.selectbox(label, prov_labels_all, index=cidx,
                                key=f"prof_{sess_key}")
            st.session_state[sess_key] = avail[prov_labels_all.index(sel)]

        if st.button("💾  Saqlash", type="primary",
                     use_container_width=True, key="save_prefs"):
            st.success("✅ Sozlamalar saqlandi!")

# ════════════════════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════════════════════

st.markdown("""
<div class="somo-footer">
    <div class="f-title">🌌 Somo AI <span class="g-text">Ultra Pro Max</span></div>
    <div style="display:flex;justify-content:center;gap:8px;flex-wrap:wrap;margin:12px 0;">
        <span class="api-badge api-groq"><span class="api-dot"></span>⚡ Groq</span>
        <span class="api-badge api-gemini"><span class="api-dot"></span>✨ Gemini</span>
        <span class="api-badge api-cohere"><span class="api-dot"></span>🔮 Cohere</span>
        <span class="api-badge api-mistral"><span class="api-dot"></span>🌪 Mistral</span>
    </div>
    <div class="f-sub">📊 Excel · 📝 Word · 💻 Kod · 🌐 HTML · 📋 CSV · 💬 Chat AI · 🔍 Tahlil</div>
    <div class="f-sub" style="margin-top:10px;">
        👨‍💻 <strong style="color:#e2e8f0;">Usmonov Sodiq</strong> · Somo AI
    </div>
    <div style="font-size:10px;color:#2a2a40;margin-top:14px;font-family:'JetBrains Mono',monospace;">
        © 2026 Somo AI Ultra Pro Max v4.0 — Python · Streamlit
    </div>
</div>
""", unsafe_allow_html=True)
